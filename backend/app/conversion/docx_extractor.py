"""Dedicated DOCX extraction helpers for the Wave Y pipeline."""

from __future__ import annotations

import base64
import html
import logging
import posixpath
import re
import zipfile
from dataclasses import dataclass, field
from io import BytesIO
from pathlib import Path
from typing import Any
from defusedxml.ElementTree import fromstring as _safe_xml_fromstring
from defusedxml.ElementTree import ParseError as _XMLParseError
from xml.etree.ElementTree import Element as _Element  # type-only; no parsing

from docx import Document
from PIL import Image

from app.conversion.archive_safety import UnsafeArchiveError, validate_ooxml_zip_archive
from app.conversion.html_generator import ir_to_html
from app.conversion.ir import IRNode

logger = logging.getLogger(__name__)

XML_NAMESPACES = {
    "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
    "cp": "http://schemas.openxmlformats.org/package/2006/metadata/core-properties",
    "dc": "http://purl.org/dc/elements/1.1/",
    "dcterms": "http://purl.org/dc/terms/",
    "o": "urn:schemas-microsoft-com:office:office",
    "pic": "http://schemas.openxmlformats.org/drawingml/2006/picture",
    "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    "rel": "http://schemas.openxmlformats.org/package/2006/relationships",
    "v": "urn:schemas-microsoft-com:vml",
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
}
WORD_DOCUMENT_PATH = "word/document.xml"
WORD_DOCUMENT_RELS_PATH = "word/_rels/document.xml.rels"
IMAGE_RELATIONSHIP_TYPE = (
    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/image"
)
IMAGE_CONTENT_TYPES = {
    ".bmp": "image/bmp",
    ".gif": "image/gif",
    ".jpeg": "image/jpeg",
    ".jpg": "image/jpeg",
    ".png": "image/png",
    ".tif": "image/tiff",
    ".tiff": "image/tiff",
    ".webp": "image/webp",
    ".wmf": "image/wmf",
}
MAX_EMBEDDED_IMAGE_SIZE = 1_000_000
COMPACT_HEADING_STYLE_RE = re.compile(r"^heading([1-6])$")
COMPACT_LIST_STYLE_RE = re.compile(r"^list(?:bullet|number)(\d+)?$")
NON_ALNUM_RE = re.compile(r"[^a-z0-9]+")
MONOSPACE_FONT_TOKENS = {
    "consolas",
    "courier",
    "couriernew",
    "lucidaconsole",
    "menlo",
    "monaco",
}


@dataclass(frozen=True, slots=True)
class ExtractionWarning:
    """Non-fatal extraction issue."""

    code: str
    message: str
    count: int | None = None


@dataclass(frozen=True, slots=True)
class HeadingItem:
    """Heading metadata for downstream TOC consumers."""

    id: str
    level: int
    text: str
    slide_number: int | None = None


@dataclass(slots=True)
class ExtractionResult:
    """Structured DOCX extraction output."""

    status: str
    html: str = ""
    title: str | None = None
    headings: list[HeadingItem] = field(default_factory=list)
    metadata: dict[str, Any] = field(default_factory=dict)
    warnings: list[ExtractionWarning] = field(default_factory=list)
    confidence: float = 0.0
    extraction_error: str | None = None
    ir: IRNode | None = None


@dataclass(frozen=True, slots=True)
class ParagraphRun:
    """Inline paragraph text segment."""

    text: str
    bold: bool = False
    italic: bool = False
    underline: bool = False
    strikethrough: bool = False
    code: bool = False


@dataclass(frozen=True, slots=True)
class StyleDefinition:
    """Resolved paragraph style metadata."""

    name: str
    list_num_id: str | None = None
    list_ilvl: int = 0


@dataclass(slots=True)
class ParagraphBlock:
    """Ordered paragraph block parsed from document.xml."""

    runs: list[ParagraphRun] = field(default_factory=list)
    style_id: str | None = None
    style_name: str | None = None
    num_id: str | None = None
    ilvl: int = 0
    is_ordered_list: bool | None = None

    @property
    def text(self) -> str:
        return "".join(run.text for run in self.runs)


@dataclass(slots=True)
class TableCellBlock:
    """Structured table cell with span metadata."""

    blocks: list["BodyBlock"] = field(default_factory=list)
    grid_column: int = 0
    colspan: int = 1
    rowspan: int = 1
    vertical_merge: str | None = None
    skip_render: bool = False


@dataclass(slots=True)
class TableRowBlock:
    """Ordered table row."""

    cells: list[TableCellBlock] = field(default_factory=list)


@dataclass(slots=True)
class TableBlock:
    """Structured table block with header/body row metadata."""

    rows: list[TableRowBlock] = field(default_factory=list)
    has_header_row: bool = False


@dataclass(frozen=True, slots=True)
class ImageBlock:
    """Embedded image rendered as a standalone figure."""

    alt: str
    src: str | None = None
    missing: bool = False


@dataclass(slots=True)
class BodyBlock:
    """Top-level or cell-level body block."""

    kind: str
    paragraph: ParagraphBlock | None = None
    table: TableBlock | None = None
    image: ImageBlock | None = None


@dataclass(slots=True)
class ParsedDocxDocument:
    """Intermediate parsed DOCX state used by later extraction tasks."""

    metadata: dict[str, Any] = field(default_factory=dict)
    style_definitions: dict[str, StyleDefinition] = field(default_factory=dict)
    numbering_definitions: dict[str, dict[int, str]] = field(default_factory=dict)
    body_blocks: list[BodyBlock] = field(default_factory=list)
    paragraphs: list[ParagraphBlock] = field(default_factory=list)
    table_count: int = 0


@dataclass(slots=True)
class ListItemBlock:
    """Single list item with optional nested lists."""

    paragraph: ParagraphBlock
    children: list["ListBlock"] = field(default_factory=list)


@dataclass(slots=True)
class ListBlock:
    """Structured list block used for nested HTML rendering."""

    ordered: bool
    num_id: str | None
    level: int
    items: list[ListItemBlock] = field(default_factory=list)


class DocxExtractor:
    """DOCX parser that validates archives and extracts ordered paragraph content."""

    def extract_path(self, file_path: str | Path) -> ExtractionResult:
        source_path = Path(file_path)
        return self.extract_bytes(source_path.read_bytes(), source_name=source_path.name)

    def extract_bytes(self, content: bytes, *, source_name: str = "<memory>") -> ExtractionResult:
        logger.info("DOCX extraction started for %s", source_name)

        try:
            with zipfile.ZipFile(BytesIO(content)) as archive:
                validate_ooxml_zip_archive(
                    archive,
                    archive_label="DOCX",
                    compressed_size_bytes=len(content),
                )
                if WORD_DOCUMENT_PATH not in archive.namelist():
                    raise ValueError("DOCX is missing word/document.xml")

                try:
                    document = Document(BytesIO(content))
                except Exception:  # policy: LOSSY — python-docx probe is advisory; XML extraction remains authoritative
                    logger.warning(
                        "DOCX python-docx probe failed for %s; continuing with XML extraction",
                        source_name,
                    )
                else:
                    logger.debug(
                        "DOCX archive opened for %s with %s python-docx paragraphs",
                        source_name,
                        len(document.paragraphs),
                    )

                parsed_document = self._parse_archive(archive)
        except zipfile.BadZipFile as exc:
            logger.exception("DOCX extraction failed for %s: invalid ZIP archive", source_name)
            return self._failed_result("Invalid DOCX archive", exc)
        except UnsafeArchiveError as exc:
            logger.warning("DOCX extraction rejected for %s: %s", source_name, exc)
            return self._failed_result("Unsafe DOCX archive", exc)
        except Exception as exc:  # policy: FAIL_FAST — extraction returns a stable failure result for unexpected parse errors
            logger.exception("DOCX extraction failed for %s", source_name)
            return self._failed_result("DOCX extraction failed", exc)

        paragraphs = self._extract_paragraphs(parsed_document)
        heading_positions = self._extract_headings(paragraphs)
        heading_lookup = {
            id(paragraphs[index]): heading
            for index, heading in heading_positions.items()
        }
        headings = list(heading_positions.values())
        table_blocks = self._extract_tables(parsed_document.body_blocks)
        image_blocks = self._extract_images(parsed_document.body_blocks)
        document_ir = self._build_ir(parsed_document.body_blocks, heading_lookup)
        html_content = ir_to_html(document_ir)
        document_title = str(parsed_document.metadata.get("title") or "") or (
            headings[0].text if headings else None
        )
        result = ExtractionResult(
            status="ready",
            html=html_content,
            title=document_title,
            headings=headings,
            metadata=parsed_document.metadata,
            warnings=[],
            confidence=1.0,
            extraction_error=None,
            ir=document_ir,
        )
        result.warnings = self._verify(result)
        result.confidence = calculate_confidence(result)

        logger.info(
            "DOCX extraction completed for %s (headings=%s, tables=%s, images=%s, confidence=%.2f)",
            source_name,
            len(result.headings),
            len(table_blocks),
            len(image_blocks),
            result.confidence,
        )
        return result

    def _parse_archive(self, archive: zipfile.ZipFile) -> ParsedDocxDocument:
        metadata = self._parse_metadata(archive)
        style_definitions = self._load_style_definitions(archive)
        numbering_definitions = self._load_numbering_definitions(archive)
        image_relationships = self._load_image_relationships(archive)
        body_blocks, paragraphs, table_count = self._parse_document_body(
            archive,
            style_definitions=style_definitions,
            numbering_definitions=numbering_definitions,
            image_relationships=image_relationships,
        )
        return ParsedDocxDocument(
            metadata=metadata,
            style_definitions=style_definitions,
            numbering_definitions=numbering_definitions,
            body_blocks=body_blocks,
            paragraphs=paragraphs,
            table_count=table_count,
        )

    def _extract_paragraphs(self, parsed_document: ParsedDocxDocument) -> list[ParagraphBlock]:
        return parsed_document.paragraphs

    def _extract_headings(self, paragraphs: list[ParagraphBlock]) -> dict[int, HeadingItem]:
        return self._build_heading_items(paragraphs)

    def _extract_lists(self, paragraphs: list[ParagraphBlock]) -> list[ListBlock]:
        return self._build_list_blocks(paragraphs)

    def _extract_tables(self, blocks: list[BodyBlock]) -> list[TableBlock]:
        tables: list[TableBlock] = []
        for block in blocks:
            if block.kind == "table" and block.table is not None:
                tables.append(block.table)
                for row in block.table.rows:
                    for cell in row.cells:
                        tables.extend(self._extract_tables(cell.blocks))
        return tables

    def _extract_images(self, blocks: list[BodyBlock]) -> list[ImageBlock]:
        images: list[ImageBlock] = []
        for block in blocks:
            if block.kind == "image" and block.image is not None:
                images.append(block.image)
            elif block.kind == "table" and block.table is not None:
                for row in block.table.rows:
                    for cell in row.cells:
                        images.extend(self._extract_images(cell.blocks))
        return images

    def _parse_metadata(self, archive: zipfile.ZipFile) -> dict[str, Any]:
        root = self._read_xml(archive, "docProps/core.xml", required=False)
        if root is None:
            return {}

        metadata: dict[str, Any] = {}
        title = self._find_text(root, "dc:title")
        author = self._find_text(root, "dc:creator")
        created = self._find_text(root, "dcterms:created")
        modified = self._find_text(root, "dcterms:modified")

        if title:
            metadata["title"] = title
        if author:
            metadata["author"] = author
        if created:
            metadata["created"] = created
        if modified:
            metadata["modified"] = modified
        return metadata

    def _load_style_definitions(self, archive: zipfile.ZipFile) -> dict[str, StyleDefinition]:
        root = self._read_xml(archive, "word/styles.xml", required=False)
        if root is None:
            return {}

        style_definitions: dict[str, StyleDefinition] = {}
        style_id_attr = self._ns_attr("w", "styleId")
        style_name_attr = self._ns_attr("w", "val")

        for style in root.findall("w:style", XML_NAMESPACES):
            style_id = style.get(style_id_attr)
            style_name = style.find("w:name", XML_NAMESPACES)
            style_label = style_name.get(style_name_attr) if style_name is not None else None
            if not style_id or not style_label:
                continue

            style_num_id, style_ilvl = self._extract_style_list_metadata(style, style_id, style_label)
            style_definitions[style_id] = StyleDefinition(
                name=style_label,
                list_num_id=style_num_id,
                list_ilvl=style_ilvl,
            )
        return style_definitions

    def _load_numbering_definitions(self, archive: zipfile.ZipFile) -> dict[str, dict[int, str]]:
        root = self._read_xml(archive, "word/numbering.xml", required=False)
        if root is None:
            return {}

        val_attr = self._ns_attr("w", "val")
        abstract_id_attr = self._ns_attr("w", "abstractNumId")
        level_attr = self._ns_attr("w", "ilvl")

        abstract_numbering: dict[str, dict[int, str]] = {}
        for abstract_num in root.findall("w:abstractNum", XML_NAMESPACES):
            abstract_id = abstract_num.get(abstract_id_attr)
            if not abstract_id:
                continue

            levels: dict[int, str] = {}
            for level in abstract_num.findall("w:lvl", XML_NAMESPACES):
                level_index = int(level.get(level_attr, "0") or "0")
                num_fmt = level.find("w:numFmt", XML_NAMESPACES)
                num_format = num_fmt.get(val_attr, "bullet") if num_fmt is not None else "bullet"
                levels[level_index] = num_format
            abstract_numbering[abstract_id] = levels

        numbering_definitions: dict[str, dict[int, str]] = {}
        num_id_attr = self._ns_attr("w", "numId")

        for num in root.findall("w:num", XML_NAMESPACES):
            num_id = num.get(num_id_attr)
            abstract_num_id = num.find("w:abstractNumId", XML_NAMESPACES)
            if not num_id or abstract_num_id is None:
                continue

            abstract_id = abstract_num_id.get(val_attr)
            if abstract_id and abstract_id in abstract_numbering:
                numbering_definitions[num_id] = abstract_numbering[abstract_id]

        return numbering_definitions

    def _load_image_relationships(self, archive: zipfile.ZipFile) -> dict[str, str]:
        root = self._read_xml(archive, WORD_DOCUMENT_RELS_PATH, required=False)
        if root is None:
            return {}

        relationships: dict[str, str] = {}
        for relationship in root.findall("rel:Relationship", XML_NAMESPACES):
            relationship_id = relationship.get("Id")
            relationship_type = relationship.get("Type")
            target = relationship.get("Target")
            target_mode = (relationship.get("TargetMode") or "").strip().lower()
            if not relationship_id or not target or relationship_type != IMAGE_RELATIONSHIP_TYPE:
                continue
            if target_mode == "external":
                continue
            relationships[relationship_id] = self._resolve_archive_target(WORD_DOCUMENT_PATH, target)
        return relationships

    def _parse_document_body(
        self,
        archive: zipfile.ZipFile,
        *,
        style_definitions: dict[str, StyleDefinition],
        numbering_definitions: dict[str, dict[int, str]],
        image_relationships: dict[str, str],
    ) -> tuple[list[BodyBlock], list[ParagraphBlock], int]:
        root = self._read_xml(archive, WORD_DOCUMENT_PATH, required=True)
        body = root.find("w:body", XML_NAMESPACES)
        if body is None:
            return [], [], 0

        body_blocks: list[BodyBlock] = []
        paragraphs: list[ParagraphBlock] = []
        table_count = 0
        image_number = 0
        for child in body:
            tag_name = self._local_name(child.tag)
            if tag_name == "p":
                paragraph = self._parse_paragraph(
                    child,
                    style_definitions=style_definitions,
                    numbering_definitions=numbering_definitions,
                )
                if paragraph.text.strip():
                    body_blocks.append(BodyBlock(kind="paragraph", paragraph=paragraph))
                    paragraphs.append(paragraph)
                image_blocks, image_number = self._extract_image_blocks(
                    child,
                    archive=archive,
                    image_relationships=image_relationships,
                    starting_image_number=image_number,
                )
                body_blocks.extend(BodyBlock(kind="image", image=image) for image in image_blocks)
                continue
            if tag_name == "tbl":
                table_block, nested_count = self._parse_table(
                    child,
                    archive=archive,
                    style_definitions=style_definitions,
                    numbering_definitions=numbering_definitions,
                    image_relationships=image_relationships,
                    starting_image_number=image_number,
                )
                body_blocks.append(BodyBlock(kind="table", table=table_block))
                table_count += 1 + nested_count
                image_number += self._count_images_in_table(table_block)

        return body_blocks, paragraphs, table_count

    def _parse_paragraph(
        self,
        paragraph_element: _Element,
        *,
        style_definitions: dict[str, StyleDefinition],
        numbering_definitions: dict[str, dict[int, str]],
    ) -> ParagraphBlock:
        paragraph_properties = paragraph_element.find("w:pPr", XML_NAMESPACES)
        style_id: str | None = None
        style_name: str | None = None
        num_id: str | None = None
        ilvl = 0
        is_ordered_list: bool | None = None

        if paragraph_properties is not None:
            style = paragraph_properties.find("w:pStyle", XML_NAMESPACES)
            style_definition: StyleDefinition | None = None
            if style is not None:
                style_id = style.get(self._ns_attr("w", "val"))
                style_definition = style_definitions.get(style_id or "")
                style_name = style_definition.name if style_definition is not None else style_id

            num_properties = paragraph_properties.find("w:numPr", XML_NAMESPACES)
            num_id, ilvl = self._resolve_paragraph_list_metadata(
                num_properties=num_properties,
                style_definition=style_definition,
            )
            if num_id is not None:
                is_ordered_list = self._resolve_list_ordering(
                    num_id=num_id,
                    ilvl=ilvl,
                    numbering_definitions=numbering_definitions,
                )

        runs = self._extract_runs(paragraph_element)
        return ParagraphBlock(
            runs=runs,
            style_id=style_id,
            style_name=style_name,
            num_id=num_id,
            ilvl=ilvl,
            is_ordered_list=is_ordered_list,
        )

    def _extract_runs(self, element: _Element) -> list[ParagraphRun]:
        runs: list[ParagraphRun] = []
        for child in element:
            tag_name = self._local_name(child.tag)
            if tag_name == "r":
                run = self._parse_run(child)
                if run is not None:
                    runs.append(run)
                continue
            if tag_name in {"hyperlink", "smartTag", "sdt", "ins"}:
                runs.extend(self._extract_runs(child))
        return runs

    def _parse_run(self, run_element: _Element) -> ParagraphRun | None:
        text = self._extract_run_text(run_element)
        if not text:
            return None

        formatting = self._extract_run_formatting(run_element)
        return ParagraphRun(text=text, **formatting)

    def _extract_run_text(self, run_element: _Element) -> str:
        parts: list[str] = []
        for node in run_element:
            tag_name = self._local_name(node.tag)
            if tag_name in {"t", "instrText"}:
                parts.append(node.text or "")
            elif tag_name == "tab":
                parts.append("\t")
            elif tag_name in {"br", "cr"}:
                parts.append("\n")
        return "".join(parts)

    def _extract_run_formatting(self, run_element: _Element) -> dict[str, bool]:
        run_properties = run_element.find("w:rPr", XML_NAMESPACES)
        if run_properties is None:
            return {
                "bold": False,
                "italic": False,
                "underline": False,
                "strikethrough": False,
                "code": False,
            }

        return {
            "bold": self._is_toggle_enabled(run_properties.find("w:b", XML_NAMESPACES)),
            "italic": self._is_toggle_enabled(run_properties.find("w:i", XML_NAMESPACES)),
            "underline": self._is_underline_enabled(run_properties.find("w:u", XML_NAMESPACES)),
            "strikethrough": self._is_toggle_enabled(
                run_properties.find("w:strike", XML_NAMESPACES)
            ),
            "code": self._is_monospace_run(run_properties),
        }

    def _build_heading_items(self, paragraphs: list[ParagraphBlock]) -> dict[int, HeadingItem]:
        headings: dict[int, HeadingItem] = {}
        seen_heading_ids: dict[str, int] = {}

        for index, paragraph in enumerate(paragraphs):
            text = paragraph.text.strip()
            if not text:
                continue

            level = self._resolve_heading_level(paragraph)
            if level is None:
                continue

            heading_id = self._build_heading_id(text, seen_heading_ids)
            headings[index] = HeadingItem(id=heading_id, level=level, text=text)

        return headings

    def _resolve_heading_level(self, paragraph: ParagraphBlock) -> int | None:
        for candidate in (paragraph.style_name, paragraph.style_id):
            normalized = self._normalize_style_token(candidate)
            if not normalized:
                continue
            if normalized == "title":
                return 1
            match = COMPACT_HEADING_STYLE_RE.fullmatch(normalized)
            if match is not None:
                return int(match.group(1))
        return None

    def _normalize_style_token(self, value: str | None) -> str:
        return NON_ALNUM_RE.sub("", (value or "").strip().lower())

    def _build_heading_id(self, text: str, seen_heading_ids: dict[str, int]) -> str:
        slug = NON_ALNUM_RE.sub("-", text.strip().lower()).strip("-") or "section"
        base_id = f"heading-{slug}"
        seen_heading_ids[base_id] = seen_heading_ids.get(base_id, 0) + 1
        duplicate_index = seen_heading_ids[base_id]
        if duplicate_index == 1:
            return base_id
        return f"{base_id}-{duplicate_index}"

    def _build_ir(
        self,
        body_blocks: list[BodyBlock],
        heading_lookup: dict[int, HeadingItem],
    ) -> IRNode:
        return IRNode(
            type="document",
            styles={"classes": ["docx-document"]},
            attributes={
                "tag": "article",
                "role": "article",
                "aria-label": "Uploaded document",
            },
            children=self._build_body_ir_blocks(body_blocks, heading_lookup=heading_lookup),
        )

    def _build_body_ir_blocks(
        self,
        blocks: list[BodyBlock],
        *,
            heading_lookup: dict[int, HeadingItem],
    ) -> list[IRNode]:
        children: list[IRNode] = []
        index = 0

        while index < len(blocks):
            block = blocks[index]
            if block.kind == "paragraph" and block.paragraph is not None:
                if self._is_list_paragraph(block.paragraph):
                    list_run, index = self._collect_list_run(blocks, index)
                    children.extend(self._build_list_ir(self._extract_lists(list_run)))
                    continue

                paragraph_node = self._build_paragraph_ir(
                    block.paragraph,
                    heading_lookup.get(id(block.paragraph)),
                )
                if paragraph_node is not None:
                    children.append(paragraph_node)
                index += 1
                continue

            if block.kind == "table" and block.table is not None:
                children.append(self._build_table_ir(block.table))
                index += 1
                continue

            if block.kind == "image" and block.image is not None:
                children.append(self._build_image_ir(block.image))
            index += 1

        return children

    def _build_paragraph_ir(
        self,
        paragraph: ParagraphBlock,
        heading: HeadingItem | None = None,
    ) -> IRNode | None:
        text = paragraph.text.strip()
        if not text:
            return None

        content = self._format_inline_runs(paragraph.runs)
        if heading is not None:
            return IRNode(
                type="heading",
                content=content,
                styles={
                    "classes": [
                        "extracted-heading",
                        f"extracted-heading-level-{heading.level}",
                    ]
                },
                attributes={"level": heading.level, "id": heading.id},
            )

        return IRNode(
            type="paragraph",
            content=content,
            styles={"classes": ["extracted-paragraph"]},
        )

    def _build_list_ir(self, list_blocks: list[ListBlock]) -> list[IRNode]:
        return [
            self._build_list_block_ir(block, root=True)
            for block in list_blocks
        ]

    def _build_list_block_ir(self, block: ListBlock, *, root: bool = False) -> IRNode:
        styles = {"classes": ["extracted-list"]} if root else {}
        return IRNode(
            type="list",
            styles=styles,
            attributes={"ordered": block.ordered},
            children=[self._build_list_item_ir(item) for item in block.items],
        )

    def _build_list_item_ir(self, item: ListItemBlock) -> IRNode:
        return IRNode(
            type="list-item",
            content=self._format_inline_runs(item.paragraph.runs),
            children=[self._build_list_block_ir(child) for child in item.children],
        )

    def _build_table_ir(self, table: TableBlock) -> IRNode:
        header_rows = table.rows[:1] if table.has_header_row and table.rows else []
        body_rows = table.rows[1:] if header_rows else table.rows
        rows = [
            self._build_table_row_ir(row, section="thead", header=True)
            for row in header_rows
        ] + [
            self._build_table_row_ir(row, section="tbody", header=False)
            for row in body_rows
        ]
        return IRNode(
            type="table",
            styles={
                "wrapper_classes": ["table-wrapper"],
                "table_classes": ["extracted-table"],
            },
            children=rows,
        )

    def _build_table_row_ir(self, row: TableRowBlock, *, section: str, header: bool) -> IRNode:
        return IRNode(
            type="table-row",
            attributes={"section": section},
            children=[
                self._build_table_cell_ir(cell, header=header)
                for cell in row.cells
                if not cell.skip_render
            ],
        )

    def _build_table_cell_ir(self, cell: TableCellBlock, *, header: bool) -> IRNode:
        attributes: dict[str, Any] = {"header": header}
        if cell.colspan > 1:
            attributes["colspan"] = cell.colspan
        if cell.rowspan > 1:
            attributes["rowspan"] = cell.rowspan

        return IRNode(
            type="table-cell",
            attributes=attributes,
            children=self._build_body_ir_blocks(cell.blocks, heading_lookup={}),
        )

    def _build_image_ir(self, image: ImageBlock) -> IRNode:
        return IRNode(
            type="image",
            styles={
                "classes": ["extracted-image"],
                "caption_classes": ["extracted-image-caption"],
                "placeholder_classes": ["extracted-image-placeholder"],
            },
            attributes={
                "src": image.src,
                "alt": image.alt,
                "caption": image.alt,
                "missing": image.missing,
                "loading": "lazy",
            },
        )

    def _format_inline_runs(self, runs: list[ParagraphRun]) -> str:
        return "".join(self._format_run(run) for run in runs)

    def _format_run(self, run: ParagraphRun) -> str:
        rendered = html.escape(run.text, quote=True).replace("\n", "<br/>")
        wrappers = (
            (run.code, "code"),
            (run.strikethrough, "del"),
            (run.underline, "u"),
            (run.italic, "em"),
            (run.bold, "strong"),
        )
        for enabled, tag in wrappers:
            if enabled:
                if tag == "code":
                    rendered = f'<code class="extracted-code">{rendered}</code>'
                else:
                    rendered = f"<{tag}>{rendered}</{tag}>"
        return rendered

    def _build_list_blocks(self, paragraphs: list[ParagraphBlock]) -> list[ListBlock]:
        blocks: list[ListBlock] = []
        index = 0
        while index < len(paragraphs):
            block, index = self._parse_list_block(
                paragraphs,
                start_index=index,
                level=paragraphs[index].ilvl,
            )
            blocks.append(block)
        return blocks

    def _parse_list_block(
        self,
        paragraphs: list[ParagraphBlock],
        *,
        start_index: int,
        level: int,
    ) -> tuple[ListBlock, int]:
        start_paragraph = paragraphs[start_index]
        block = ListBlock(
            ordered=bool(start_paragraph.is_ordered_list),
            num_id=start_paragraph.num_id,
            level=level,
        )
        index = start_index

        while index < len(paragraphs):
            paragraph = paragraphs[index]
            if paragraph.ilvl < level:
                break

            if paragraph.ilvl > level:
                if not block.items:
                    break
                child_block, index = self._parse_list_block(
                    paragraphs,
                    start_index=index,
                    level=paragraph.ilvl,
                )
                block.items[-1].children.append(child_block)
                continue

            if block.items and (
                paragraph.num_id != block.num_id
                or bool(paragraph.is_ordered_list) != block.ordered
            ):
                break

            block.items.append(ListItemBlock(paragraph=paragraph))
            index += 1

        return block, index

    def _collect_list_run(
        self,
        blocks: list[BodyBlock],
        start_index: int,
    ) -> tuple[list[ParagraphBlock], int]:
        index = start_index
        run: list[ParagraphBlock] = []
        while index < len(blocks):
            block = blocks[index]
            if block.kind != "paragraph" or block.paragraph is None:
                break
            if not self._is_list_paragraph(block.paragraph):
                break
            run.append(block.paragraph)
            index += 1
        return run, index

    def _is_list_paragraph(self, paragraph: ParagraphBlock) -> bool:
        return paragraph.num_id is not None

    def _parse_table(
        self,
        table_element: _Element,
        *,
        archive: zipfile.ZipFile,
        style_definitions: dict[str, StyleDefinition],
        numbering_definitions: dict[str, dict[int, str]],
        image_relationships: dict[str, str],
        starting_image_number: int,
    ) -> tuple[TableBlock, int]:
        table = TableBlock(has_header_row=self._table_has_header_row(table_element))
        nested_table_count = 0
        image_number = starting_image_number

        for row_element in table_element.findall("w:tr", XML_NAMESPACES):
            row = TableRowBlock()
            grid_column = 0
            for cell_element in row_element.findall("w:tc", XML_NAMESPACES):
                cell, child_table_count = self._parse_table_cell(
                    cell_element,
                    archive=archive,
                    grid_column=grid_column,
                    style_definitions=style_definitions,
                    numbering_definitions=numbering_definitions,
                    image_relationships=image_relationships,
                    starting_image_number=image_number,
                )
                row.cells.append(cell)
                grid_column += cell.colspan
                nested_table_count += child_table_count
                image_number += self._count_images_in_blocks(cell.blocks)
            table.rows.append(row)

        self._apply_rowspans(table)
        return table, nested_table_count

    def _parse_table_cell(
        self,
        cell_element: _Element,
        *,
        archive: zipfile.ZipFile,
        grid_column: int,
        style_definitions: dict[str, StyleDefinition],
        numbering_definitions: dict[str, dict[int, str]],
        image_relationships: dict[str, str],
        starting_image_number: int,
    ) -> tuple[TableCellBlock, int]:
        cell_properties = cell_element.find("w:tcPr", XML_NAMESPACES)
        colspan = 1
        vertical_merge: str | None = None
        if cell_properties is not None:
            grid_span = cell_properties.find("w:gridSpan", XML_NAMESPACES)
            if grid_span is not None:
                colspan = int(grid_span.get(self._ns_attr("w", "val"), "1") or "1")

            vmerge = cell_properties.find("w:vMerge", XML_NAMESPACES)
            if vmerge is not None:
                vertical_merge = (vmerge.get(self._ns_attr("w", "val")) or "continue").strip()

        blocks, nested_table_count = self._parse_container_blocks(
            cell_element,
            archive=archive,
            style_definitions=style_definitions,
            numbering_definitions=numbering_definitions,
            image_relationships=image_relationships,
            starting_image_number=starting_image_number,
        )
        return (
            TableCellBlock(
                blocks=blocks,
                grid_column=grid_column,
                colspan=colspan,
                vertical_merge=vertical_merge,
            ),
            nested_table_count,
        )

    def _parse_container_blocks(
        self,
        container: _Element,
        *,
        archive: zipfile.ZipFile,
        style_definitions: dict[str, StyleDefinition],
        numbering_definitions: dict[str, dict[int, str]],
        image_relationships: dict[str, str],
        starting_image_number: int,
    ) -> tuple[list[BodyBlock], int]:
        blocks: list[BodyBlock] = []
        nested_table_count = 0
        image_number = starting_image_number

        for child in container:
            tag_name = self._local_name(child.tag)
            if tag_name == "p":
                paragraph = self._parse_paragraph(
                    child,
                    style_definitions=style_definitions,
                    numbering_definitions=numbering_definitions,
                )
                if paragraph.text.strip():
                    blocks.append(BodyBlock(kind="paragraph", paragraph=paragraph))
                image_blocks, image_number = self._extract_image_blocks(
                    child,
                    archive=archive,
                    image_relationships=image_relationships,
                    starting_image_number=image_number,
                )
                blocks.extend(BodyBlock(kind="image", image=image) for image in image_blocks)
                continue

            if tag_name == "tbl":
                table, child_nested_count = self._parse_table(
                    child,
                    archive=archive,
                    style_definitions=style_definitions,
                    numbering_definitions=numbering_definitions,
                    image_relationships=image_relationships,
                    starting_image_number=image_number,
                )
                blocks.append(BodyBlock(kind="table", table=table))
                nested_table_count += 1 + child_nested_count
                image_number += self._count_images_in_table(table)

        return blocks, nested_table_count

    def _table_has_header_row(self, table_element: _Element) -> bool:
        table_look = table_element.find("w:tblPr/w:tblLook", XML_NAMESPACES)
        if table_look is not None:
            first_row = (table_look.get(self._ns_attr("w", "firstRow")) or "").strip().lower()
            if first_row not in {"", "0", "false", "off"}:
                return True

        first_row_element = table_element.find("w:tr", XML_NAMESPACES)
        if first_row_element is None:
            return False
        return first_row_element.find("w:trPr/w:tblHeader", XML_NAMESPACES) is not None

    def _apply_rowspans(self, table: TableBlock) -> None:
        for row_index, row in enumerate(table.rows):
            for cell in row.cells:
                if cell.vertical_merge == "continue":
                    cell.skip_render = True
                    continue
                if cell.vertical_merge != "restart":
                    continue

                rowspan = 1
                next_row_index = row_index + 1
                while next_row_index < len(table.rows):
                    continuation = self._find_table_cell(
                        table.rows[next_row_index],
                        grid_column=cell.grid_column,
                    )
                    if continuation is None or continuation.vertical_merge != "continue":
                        break
                    continuation.skip_render = True
                    rowspan += 1
                    next_row_index += 1
                cell.rowspan = rowspan

    def _find_table_cell(self, row: TableRowBlock, *, grid_column: int) -> TableCellBlock | None:
        for cell in row.cells:
            if cell.grid_column == grid_column:
                return cell
        return None

    def _verify(self, result: ExtractionResult) -> list[ExtractionWarning]:
        warnings: list[ExtractionWarning] = []
        if not self._has_visible_text(result.html):
            warnings.append(ExtractionWarning(code="NO_CONTENT", message="Document is empty"))
        missing_image_count = result.html.count("[Image")
        if missing_image_count > 0:
            warnings.append(
                ExtractionWarning(
                    code="MISSING_IMAGES",
                    message=f"{missing_image_count} images failed",
                    count=missing_image_count,
                )
            )
        if "<table" in result.html and not self._valid_tables(result.html):
            warnings.append(ExtractionWarning(code="BAD_TABLES", message="Tables may be malformed"))
        return warnings

    def _has_visible_text(self, html_content: str) -> bool:
        if "<img" in (html_content or ""):
            return True
        # M-30: Use proper XML parser instead of fragile regex for tag stripping
        try:
            root = _safe_xml_fromstring(f"<r>{html_content or ''}</r>")
            text = " ".join(root.itertext()).strip()
            return bool(text)
        except _XMLParseError:
            return bool((html_content or "").strip())

    def _valid_tables(self, html_content: str) -> bool:
        try:
            root = _safe_xml_fromstring(f"<root>{html_content}</root>")
        except _XMLParseError:
            return False

        for table in root.findall(".//table"):
            if table.find("thead") is None or table.find("tbody") is None:
                return False
        return True

    def _read_xml(
        self,
        archive: zipfile.ZipFile,
        archive_path: str,
        *,
        required: bool,
    ) -> _Element | None:
        try:
            xml_bytes = archive.read(archive_path)
        except KeyError:
            if required:
                raise ValueError(f"DOCX archive is missing {archive_path}") from None
            return None

        try:
            return _safe_xml_fromstring(xml_bytes)
        except _XMLParseError:
            if required:
                raise ValueError(f"DOCX XML at {archive_path} could not be parsed") from None
            logger.warning("Skipping unreadable optional DOCX XML: %s", archive_path)
            return None

    def _find_text(self, root: _Element, path: str) -> str | None:
        element = root.find(path, XML_NAMESPACES)
        if element is None or element.text is None:
            return None
        text = element.text.strip()
        return text or None

    def _failed_result(self, message: str, exc: Exception) -> ExtractionResult:
        warning = ExtractionWarning(code="PARSE_FAILED", message=f"{message}: {exc}")
        return ExtractionResult(
            status="failed",
            html="",
            title=None,
            headings=[],
            metadata={},
            warnings=[warning],
            confidence=0.0,
            extraction_error=f"{message}: {exc}",
        )

    def _local_name(self, tag: str) -> str:
        return tag.rsplit("}", maxsplit=1)[-1]

    def _ns_attr(self, prefix: str, attribute: str) -> str:
        return f"{{{XML_NAMESPACES[prefix]}}}{attribute}"

    def _is_toggle_enabled(self, element: _Element | None) -> bool:
        if element is None:
            return False
        value = (element.get(self._ns_attr("w", "val")) or "").strip().lower()
        return value not in {"0", "false", "off"}

    def _is_underline_enabled(self, element: _Element | None) -> bool:
        if element is None:
            return False
        value = (element.get(self._ns_attr("w", "val")) or "single").strip().lower()
        return value not in {"", "0", "false", "none", "off"}

    def _is_monospace_run(self, run_properties: _Element) -> bool:
        run_fonts = run_properties.find("w:rFonts", XML_NAMESPACES)
        if run_fonts is None:
            return False

        font_attributes = ("ascii", "hAnsi", "cs", "eastAsia")
        for attribute_name in font_attributes:
            font_name = run_fonts.get(self._ns_attr("w", attribute_name))
            normalized = self._normalize_style_token(font_name)
            if normalized and normalized in MONOSPACE_FONT_TOKENS:
                return True
        return False

    def _extract_image_blocks(
        self,
        paragraph_element: _Element,
        *,
        archive: zipfile.ZipFile,
        image_relationships: dict[str, str],
        starting_image_number: int,
    ) -> tuple[list[ImageBlock], int]:
        image_blocks: list[ImageBlock] = []
        image_number = starting_image_number

        for node in paragraph_element.iter():
            is_image_node, relationship_id = self._extract_supported_image_relationship(node)
            if not is_image_node:
                continue
            image_number += 1
            image_blocks.append(
                self._parse_image_block(
                    relationship_id=relationship_id,
                    archive=archive,
                    image_relationships=image_relationships,
                    image_number=image_number,
                )
            )

        return image_blocks, image_number

    def _parse_image_block(
        self,
        *,
        relationship_id: str | None,
        archive: zipfile.ZipFile,
        image_relationships: dict[str, str],
        image_number: int,
    ) -> ImageBlock:
        alt = f"Figure {image_number}"
        if not relationship_id:
            return ImageBlock(alt=alt, missing=True)

        archive_path = image_relationships.get(relationship_id)
        if not archive_path:
            return ImageBlock(alt=alt, missing=True)

        try:
            image_bytes = archive.read(archive_path)
        except KeyError:
            return ImageBlock(alt=alt, missing=True)

        content_type = self._guess_image_content_type(archive_path)
        prepared_bytes, prepared_content_type = self._prepare_image_for_embedding(
            image_bytes,
            content_type=content_type,
        )
        data_url = (
            f"data:{prepared_content_type};base64,"
            f"{base64.b64encode(prepared_bytes).decode('ascii')}"
        )
        return ImageBlock(alt=alt, src=data_url, missing=False)

    def _extract_supported_image_relationship(
        self,
        image_element: _Element,
    ) -> tuple[bool, str | None]:
        local_name = self._local_name(image_element.tag)
        if local_name == "drawing":
            return self._extract_drawing_relationship_id(image_element)
        if local_name == "pict":
            return self._extract_pict_relationship_id(image_element)
        return False, None

    def _extract_drawing_relationship_id(
        self,
        drawing_element: _Element,
    ) -> tuple[bool, str | None]:
        blip = drawing_element.find(".//a:blip", XML_NAMESPACES)
        if blip is None:
            return False, None

        for attribute_name in ("embed", "link"):
            relationship_id = blip.get(self._ns_attr("r", attribute_name))
            if relationship_id:
                return True, relationship_id
        return True, None

    def _extract_pict_relationship_id(
        self,
        pict_element: _Element,
    ) -> tuple[bool, str | None]:
        image_data = pict_element.find(".//v:imagedata", XML_NAMESPACES)
        if image_data is None:
            return False, None

        relationship_candidates = (
            image_data.get(self._ns_attr("r", "id")),
            image_data.get(self._ns_attr("o", "relid")),
            image_data.get("id"),
            image_data.get("relid"),
        )
        for relationship_id in relationship_candidates:
            if relationship_id:
                return True, relationship_id
        return True, None

    def _resolve_archive_target(self, source_part: str, target: str) -> str:
        source_dir = posixpath.dirname(source_part)
        return posixpath.normpath(posixpath.join(source_dir, target))

    def _guess_image_content_type(self, archive_path: str) -> str:
        extension = Path(archive_path).suffix.lower()
        return IMAGE_CONTENT_TYPES.get(extension, "image/png")

    def _prepare_image_for_embedding(
        self,
        image_bytes: bytes,
        *,
        content_type: str,
    ) -> tuple[bytes, str]:
        if len(image_bytes) <= MAX_EMBEDDED_IMAGE_SIZE:
            return image_bytes, content_type

        try:
            return compress_image_for_embedding(
                image_bytes,
                max_size=MAX_EMBEDDED_IMAGE_SIZE,
            )
        except Exception:  # policy: LOSSY — image compression fallback should preserve extraction output
            logger.warning("Falling back to original image bytes after compression failure")
            return image_bytes, content_type

    def _count_images_in_blocks(self, blocks: list[BodyBlock]) -> int:
        return sum(1 for block in blocks if block.kind == "image" and block.image is not None)

    def _count_images_in_table(self, table: TableBlock) -> int:
        return sum(self._count_images_in_blocks(cell.blocks) for row in table.rows for cell in row.cells)

    def _extract_style_list_metadata(
        self,
        style_element: _Element,
        style_id: str,
        style_label: str,
    ) -> tuple[str | None, int]:
        paragraph_properties = style_element.find("w:pPr", XML_NAMESPACES)
        if paragraph_properties is None:
            return None, 0

        num_properties = paragraph_properties.find("w:numPr", XML_NAMESPACES)
        if num_properties is None:
            return None, 0

        num_id_element = num_properties.find("w:numId", XML_NAMESPACES)
        if num_id_element is None:
            return None, 0

        num_id = num_id_element.get(self._ns_attr("w", "val"))
        ilvl_element = num_properties.find("w:ilvl", XML_NAMESPACES)
        if ilvl_element is not None:
            ilvl = int(ilvl_element.get(self._ns_attr("w", "val"), "0") or "0")
            return num_id, ilvl

        inferred_level = self._infer_style_list_level(style_id, style_label)
        return num_id, inferred_level

    def _infer_style_list_level(self, style_id: str, style_label: str) -> int:
        for candidate in (style_id, style_label):
            normalized = self._normalize_style_token(candidate)
            match = COMPACT_LIST_STYLE_RE.fullmatch(normalized)
            if match is None:
                continue
            if match.group(1) is None:
                return 0
            return max(0, int(match.group(1)) - 1)
        return 0

    def _resolve_paragraph_list_metadata(
        self,
        *,
        num_properties: _Element | None,
        style_definition: StyleDefinition | None,
    ) -> tuple[str | None, int]:
        if num_properties is not None:
            num_id_element = num_properties.find("w:numId", XML_NAMESPACES)
            ilvl_element = num_properties.find("w:ilvl", XML_NAMESPACES)
            num_id = (
                num_id_element.get(self._ns_attr("w", "val"))
                if num_id_element is not None
                else None
            )
            ilvl = (
                int(ilvl_element.get(self._ns_attr("w", "val"), "0") or "0")
                if ilvl_element is not None
                else (style_definition.list_ilvl if style_definition is not None else 0)
            )
            return num_id, ilvl

        if style_definition is None or style_definition.list_num_id is None:
            return None, 0
        return style_definition.list_num_id, style_definition.list_ilvl

    def _resolve_list_ordering(
        self,
        *,
        num_id: str,
        ilvl: int,
        numbering_definitions: dict[str, dict[int, str]],
    ) -> bool:
        level_definitions = numbering_definitions.get(num_id, {})
        num_format = level_definitions.get(ilvl) or level_definitions.get(0) or "bullet"
        return num_format != "bullet"


def calculate_confidence(result: ExtractionResult) -> float:
    """Best-effort confidence score based on current warning set."""

    score = 1.0
    for warning in result.warnings:
        if warning.code == "PARSE_FAILED":
            score -= 1.0
        elif warning.code == "MISSING_IMAGES":
            score -= 0.1 * float(warning.count or 1)
        elif warning.code == "BAD_TABLES":
            score -= 0.2
        elif warning.code == "NO_CONTENT":
            score -= 0.5
    return max(0.0, min(1.0, score))


def compress_image_for_embedding(
    image_bytes: bytes,
    max_size: int = MAX_EMBEDDED_IMAGE_SIZE,
) -> tuple[bytes, str]:
    """Compress image for base64 embedding."""

    image = Image.open(BytesIO(image_bytes))
    if image.mode in {"RGBA", "LA"}:
        image = image.convert("RGB")

    max_dimension = 1200
    if max(image.size) > max_dimension:
        image.thumbnail((max_dimension, max_dimension), Image.Resampling.LANCZOS)

    output = BytesIO()
    quality = 85
    while quality > 20:
        output.seek(0)
        output.truncate()
        image.save(output, "JPEG", quality=quality, optimize=True)
        if output.tell() <= max_size:
            break
        quality -= 10

    return output.getvalue(), "image/jpeg"


def extract_docx(file_path: str | Path) -> ExtractionResult:
    """Parse a DOCX file path into a structured extraction result."""

    return DocxExtractor().extract_path(file_path)


__all__ = [
    "DocxExtractor",
    "ExtractionResult",
    "ExtractionWarning",
    "HeadingItem",
    "calculate_confidence",
    "compress_image_for_embedding",
    "extract_docx",
]
