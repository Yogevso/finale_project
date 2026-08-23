"""PDF-to-DOCX converter.

Extracts text, tables, images, and basic structure from a PDF using PyMuPDF
(fitz), then builds a DOCX file via python-docx.  The resulting DOCX bytes are
fed through the existing DocxExtractor pipeline so upstreamed PDFs behave
*exactly* like natively uploaded DOCX files.

Structure recovery notes:

- Tables are detected with ``page.find_tables()`` and emitted as real DOCX
  tables.  Text living inside a detected table is claimed by that table so it
  is not also emitted as loose paragraphs.
- Paragraphs are assembled per *block*, not per line, so a wrapped sentence
  stays one paragraph instead of one paragraph per visual line.
- Headings are detected relative to the document's dominant body font size,
  because an absolute point threshold misfires on cover pages and on symbol
  glyphs.
- Running headers/footers and watermarks repeat on most pages; they are
  detected in a first pass and dropped so they do not pollute the reader view.
"""

from __future__ import annotations

import io
import logging
import re
from collections import Counter
from dataclasses import dataclass, field
from typing import Any

import fitz  # PyMuPDF
from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

logger = logging.getLogger(__name__)

# Image dimension cap inside the generated DOCX (inches).
_MAX_IMG_WIDTH_IN = 6.0
_MAX_IMG_HEIGHT_IN = 8.0
# Minimum text length to consider a block meaningful.
_MIN_BLOCK_TEXT_LEN = 1
# Heading thresholds, expressed as a multiple of the dominant body font size.
_HEADING_RATIO = 1.45
_SUBHEADING_RATIO = 1.15
# A heading candidate shorter than this is almost always a stray glyph or a
# superscript mark rather than a real heading.
_MIN_HEADING_CHARS = 3
# Longer than this and the run is prose that happens to be set large.
_MAX_HEADING_CHARS = 120
# Fallback body size when a PDF exposes no usable span sizes.
_DEFAULT_BODY_FONT_PT = 11.0
# A run of text repeating on at least this fraction of pages is chrome
# (running header/footer or watermark), not content.
_REPEAT_CHROME_RATIO = 0.6
# Minimum pages before repeat-detection is trustworthy.
_MIN_PAGES_FOR_REPEAT_DETECTION = 4
# An image reused on at least this fraction of pages is a logo/watermark.
_REPEAT_IMAGE_RATIO = 0.5
# Quality for images re-encoded as JPEG. Across a twenty-page Intel guide the same
# 42 screenshots weigh 8.97 MB as PNG and 1.60 MB at this setting, against 1.44 MB
# for the streams the PDF itself stores.
_JPEG_QUALITY = 85
# Past this, a document's images are worth saying out loud. The reader artifact and
# every version of the document carry them base64-encoded, a third larger again, and
# the largest version stored today is half a megabyte. Chosen to sit above an ordinary
# illustrated guide - the twenty-page Intel one lands at 1.6 MB - and below the point
# where one upload is heavier than the whole versions table.
_IMAGE_BUDGET_BYTES = 6 * 1024 * 1024
# Images smaller than this (points, roughly 1/72") carry no reader value.
_MIN_IMAGE_DIMENSION = 32
# Fraction of a text block that must sit inside a table to be claimed by it.
_TABLE_CLAIM_OVERLAP = 0.5
# A single cell holding at least this share of a table's text means the grid was
# mis-detected and the row collapsed instead of splitting across its columns.
_COLLAPSED_CELL_SHARE = 0.4
# Below this many rows a single filled cell is an ordinary title row, not a failure.
_MIN_ROWS_FOR_COLLAPSE_CHECK = 3
# A collapsed row may only be dropped when every one of its words already appears
# elsewhere in the table. Anything short of that would lose content.
# Optional-content layer names that mark decorative page furniture.
_DECORATIVE_OCG_NAMES = {"background", "watermark", "watermarks"}

# Symbol/Wingdings bullets land in the Unicode private use area.
_PUA_RANGE = (0xF000, 0xF0FF)
_BULLET_CHAR = "•"


@dataclass
class _ExtractedBlock:
    """Intermediate representation of a content block from a PDF page."""

    kind: str  # "text" | "image" | "table"
    text: str = ""
    font_size: float = 0.0
    is_bold: bool = False
    is_bullet: bool = False
    image_bytes: bytes = b""
    image_ext: str = "png"
    table_rows: list[list[str]] = field(default_factory=list)
    order: tuple[float, float] = (0.0, 0.0)


@dataclass
class PdfConversionResult:
    """Result of converting a PDF to DOCX bytes."""

    docx_bytes: bytes = b""
    page_count: int = 0
    warnings: list[str] = field(default_factory=list)
    error: str | None = None


def convert_pdf_to_docx(pdf_bytes: bytes) -> PdfConversionResult:
    """Convert raw PDF bytes into DOCX bytes.

    The generated DOCX preserves:
    - Text paragraphs with body-relative heading detection
    - Tables, as real DOCX tables
    - Embedded images (raster only, capped at page width), minus repeated logos
    - Page breaks between PDF pages

    Returns a ``PdfConversionResult`` with the DOCX bytes or an error.
    """
    result = PdfConversionResult()
    try:
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    except (
        Exception
    ) as exc:  # policy: FAIL_FAST — invalid PDF input returns a stable conversion error
        result.error = f"Failed to open PDF: {exc}"
        return result

    result.page_count = len(doc)
    if result.page_count == 0:
        result.error = "PDF has no pages"
        doc.close()
        return result

    ocr_enabled = _is_ocr_enabled()
    body_font_size = _detect_body_font_size(doc)
    chrome_text = _detect_repeated_text(doc)
    repeated_xrefs = _detect_repeated_images(doc)

    blocks_by_page: list[list[_ExtractedBlock]] = []

    for page_idx in range(len(doc)):
        page = doc[page_idx]
        page_blocks: list[_ExtractedBlock] = []

        table_regions, table_blocks = _extract_tables(page, page_idx, result.warnings)
        page_blocks.extend(table_blocks)

        text_blocks, text_found = _extract_text_blocks(
            page,
            body_font_size=body_font_size,
            chrome_text=chrome_text,
            table_regions=table_regions,
        )
        page_blocks.extend(text_blocks)

        # AH-004: OCR fallback for image-only pages (scanned PDFs)
        if not text_found and not table_blocks and ocr_enabled:
            ocr_text = _ocr_page(page, page_idx, result.warnings)
            if ocr_text:
                page_blocks.append(
                    _ExtractedBlock(
                        kind="text",
                        text=ocr_text,
                        font_size=body_font_size,
                        order=(0.0, 0.0),
                    )
                )

        page_blocks.extend(_extract_images(page, page_idx, doc, repeated_xrefs, result.warnings))

        page_blocks.sort(key=lambda blk: blk.order)
        blocks_by_page.append(page_blocks)

    doc.close()

    # --- build DOCX ---
    docx_doc = DocxDocument()
    style = docx_doc.styles["Normal"]
    style.font.size = Pt(11)

    for page_idx, page_blocks in enumerate(blocks_by_page):
        if page_idx > 0:
            docx_doc.add_page_break()

        for blk in page_blocks:
            if blk.kind == "text":
                _add_text_block(docx_doc, blk, body_font_size)
            elif blk.kind == "table":
                _add_table_block(docx_doc, blk, result.warnings, page_idx)
            elif blk.kind == "image":
                _add_image_block(docx_doc, blk, result.warnings, page_idx)

    _warn_on_image_weight(blocks_by_page, result.warnings)

    buf = io.BytesIO()
    docx_doc.save(buf)
    result.docx_bytes = buf.getvalue()
    return result


def _warn_on_image_weight(
    blocks_by_page: list[list[_ExtractedBlock]],
    warnings: list[str],
) -> None:
    """Say how heavy a document's images are before they are anyone's problem.

    Nothing is downscaled here on purpose. These are screenshots of interfaces, and the
    text inside them is the reason the reader opened the page - a cap that quietly
    softens them trades the thing the document is for against bytes. Measured on the
    twenty-page Intel guide, capping the long edge at 1600px saves 24% while the choice
    of JPEG over PNG already saved 82%, so the cap buys little and costs legibility.
    """
    total = sum(
        len(blk.image_bytes)
        for page_blocks in blocks_by_page
        for blk in page_blocks
        if blk.kind == "image"
    )
    if total > _IMAGE_BUDGET_BYTES:
        warnings.append(
            f"This document's images weigh {total / 1024 / 1024:.1f} MB, which every "
            "stored version of it will carry. Consider whether they all belong in the "
            "document body."
        )


# ---------------------------------------------------------------------------
# Document-level analysis passes
# ---------------------------------------------------------------------------


def _detect_body_font_size(doc: Any) -> float:
    """Return the dominant body font size, weighted by how much text uses it.

    Heading detection keys off this instead of an absolute point value, because
    the same 15pt run is a heading in a 9pt datasheet and body text on a cover
    page.
    """
    weights: Counter[float] = Counter()
    for page_idx in range(min(len(doc), 20)):
        try:
            text_dict = doc[page_idx].get_text("dict")
        except Exception:  # policy: DEGRADED — a page that will not parse cannot inform sizing
            continue
        for block in text_dict.get("blocks", []):
            if block.get("type") != 0:
                continue
            for line in block.get("lines", []):
                for span in line.get("spans", []):
                    text = span.get("text", "").strip()
                    if text:
                        weights[round(span.get("size", 0.0), 1)] += len(text)
    if not weights:
        return _DEFAULT_BODY_FONT_PT
    return max(weights.items(), key=lambda item: item[1])[0] or _DEFAULT_BODY_FONT_PT


def _normalize_chrome_text(text: str) -> str:
    """Normalize a line so page-varying numbers do not defeat repeat matching."""
    return re.sub(r"\d+", "#", text.strip().lower())


def _detect_repeated_text(doc: Any) -> set[str]:
    """Return normalized text that repeats on most pages (headers/footers/watermarks)."""
    page_count = len(doc)
    if page_count < _MIN_PAGES_FOR_REPEAT_DETECTION:
        return set()

    seen: Counter[str] = Counter()
    sampled = 0
    for page_idx in range(min(page_count, 40)):
        try:
            page = doc[page_idx]
            text_dict = page.get_text("dict")
        except Exception:  # policy: DEGRADED — unreadable page simply does not vote
            continue
        sampled += 1
        page_lines: set[str] = set()
        for block in text_dict.get("blocks", []):
            if block.get("type") != 0:
                continue
            for line in block.get("lines", []):
                text = "".join(span.get("text", "") for span in line.get("spans", []))
                normalized = _normalize_chrome_text(text)
                if normalized:
                    page_lines.add(normalized)
        seen.update(page_lines)

    if sampled < _MIN_PAGES_FOR_REPEAT_DETECTION:
        return set()
    threshold = max(2, int(sampled * _REPEAT_CHROME_RATIO))
    return {text for text, count in seen.items() if count >= threshold}


def _detect_repeated_images(doc: Any) -> set[int]:
    """Return xrefs of images reused across most pages (logos and watermarks)."""
    page_count = len(doc)
    counts: Counter[int] = Counter()
    sampled = 0
    for page_idx in range(min(page_count, 40)):
        try:
            images = doc[page_idx].get_images(full=True)
        except Exception:  # policy: DEGRADED — unreadable page simply does not vote
            continue
        sampled += 1
        for xref in {info[0] for info in images}:
            counts[xref] += 1
    if sampled < 2:
        return set()
    threshold = max(2, int(sampled * _REPEAT_IMAGE_RATIO))
    return {xref for xref, count in counts.items() if count >= threshold}


# ---------------------------------------------------------------------------
# Per-page extraction
# ---------------------------------------------------------------------------


def _extract_tables(
    page: Any,
    page_idx: int,
    warnings: list[str],
) -> tuple[list[Any], list[_ExtractedBlock]]:
    """Detect tables on a page and return their regions plus emitted blocks."""
    regions: list[Any] = []
    blocks: list[_ExtractedBlock] = []
    try:
        found = page.find_tables()
    except Exception as exc:  # policy: LOSSY — table detection failure degrades to plain text
        warnings.append(f"Page {page_idx + 1}: table detection failed: {exc}")
        return regions, blocks

    for table in getattr(found, "tables", []) or []:
        try:
            rows = table.extract()
        except Exception as exc:  # policy: LOSSY — one bad table must not drop the page
            warnings.append(f"Page {page_idx + 1}: table extraction failed: {exc}")
            continue
        cleaned = _clean_table_rows(rows)
        if not cleaned:
            continue
        try:
            rect = fitz.Rect(table.bbox)
        except Exception:  # policy: DEGRADED — without a bbox the table cannot claim text
            rect = None
        cleaned = _drop_empty_columns(_repair_collapsed_rows(cleaned, page_idx, warnings))
        if len(cleaned) < 2:
            continue
        if rect is not None:
            regions.append(rect)
        blocks.append(
            _ExtractedBlock(
                kind="table",
                table_rows=cleaned,
                order=(rect.y0, rect.x0) if rect is not None else (0.0, 0.0),
            )
        )
    return regions, blocks


def _row_is_collapsed(row: list[str], table_text_len: int) -> bool:
    """True when one cell swallowed the row instead of the text splitting by column."""
    filled = [cell for cell in row if cell]
    if len(row) < 2 or len(filled) != 1 or table_text_len <= 0:
        return False
    return len(filled[0]) >= table_text_len * _COLLAPSED_CELL_SHARE


def _has_collapsed_row(rows: list[list[str]]) -> bool:
    if len(rows) < _MIN_ROWS_FOR_COLLAPSE_CHECK:
        return False
    total = sum(len(cell) for row in rows for cell in row)
    return any(_row_is_collapsed(row, total) for row in rows)


def _table_words(text: str) -> list[str]:
    return re.findall(r"\w+", text.lower())


def _is_covered_elsewhere(cell: str, other_text: str) -> bool:
    """True when a cell carries no word the rest of the table does not already hold."""
    available = set(_table_words(other_text))
    return all(word in available for word in _table_words(cell))


def _repair_collapsed_rows(
    rows: list[list[str]],
    page_idx: int,
    warnings: list[str],
) -> list[list[str]]:
    """Drop rows that collapsed into one cell, but only where nothing is lost."""
    if len(rows) < _MIN_ROWS_FOR_COLLAPSE_CHECK:
        return rows

    table_text_len = sum(len(cell) for row in rows for cell in row)
    kept: list[list[str]] = []
    for index, row in enumerate(rows):
        if not _row_is_collapsed(row, table_text_len):
            kept.append(row)
            continue

        collapsed_cell = next(cell for cell in row if cell)
        elsewhere = " ".join(
            cell for other, values in enumerate(rows) if other != index for cell in values
        )
        if not _is_covered_elsewhere(collapsed_cell, elsewhere):
            # Dropping this would remove text the table does not carry anywhere else.
            kept.append(row)
            continue

        warnings.append(
            f"Page {page_idx + 1}: dropped a mis-detected table row that repeated the table."
        )

    return kept if len(kept) >= 2 else rows


def _drop_empty_columns(rows: list[list[str]]) -> list[list[str]]:
    """Remove phantom columns that mis-detection inserts and that hold no text."""
    if not rows:
        return rows

    width = max(len(row) for row in rows)
    keep = [index for index in range(width) if any(_cell_at(row, index) for row in rows)]
    if len(keep) == width or len(keep) < 2:
        return rows
    return [[_cell_at(row, index) for index in keep] for row in rows]


def _cell_at(row: list[str], index: int) -> str:
    return row[index] if index < len(row) else ""


def _clean_table_rows(rows: list[list[str | None]]) -> list[list[str]]:
    """Normalize extracted cells and drop tables with no real content."""
    cleaned: list[list[str]] = []
    for row in rows or []:
        cells = [_normalize_cell(cell) for cell in row]
        if any(cell for cell in cells):
            cleaned.append(cells)
    if len(cleaned) < 2:
        return []
    width = max(len(row) for row in cleaned)
    if width < 2:
        return []
    return [row + [""] * (width - len(row)) for row in cleaned]


def _normalize_cell(cell: str | None) -> str:
    """Collapse the newlines PyMuPDF inserts at cell line wraps."""
    if not cell:
        return ""
    return re.sub(r"\s+", " ", _demote_pua(str(cell))).strip()


def _demote_pua(text: str) -> str:
    """Replace symbol-font private-use glyphs with a plain bullet."""
    return "".join(_BULLET_CHAR if _PUA_RANGE[0] <= ord(ch) <= _PUA_RANGE[1] else ch for ch in text)


def _rect_overlap_ratio(inner: Any, outer: Any) -> float:
    """Fraction of ``inner`` that lies inside ``outer``."""
    try:
        intersection = fitz.Rect(inner) & fitz.Rect(outer)
    except Exception:  # policy: DEGRADED — unusable geometry means no overlap
        return 0.0
    area = abs(fitz.Rect(inner).get_area())
    if area <= 0 or intersection.is_empty:
        return 0.0
    return abs(intersection.get_area()) / area


def _extract_text_blocks(
    page: Any,
    *,
    body_font_size: float,
    chrome_text: set[str],
    table_regions: list[Any],
) -> tuple[list[_ExtractedBlock], bool]:
    """Extract paragraph-level text blocks, skipping table interiors and chrome."""
    blocks: list[_ExtractedBlock] = []
    text_found = False
    try:
        text_dict = page.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE, sort=True)
    except TypeError:
        # Older PyMuPDF builds do not accept ``sort`` on get_text().
        text_dict = page.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE)
    except Exception:  # policy: LOSSY — an unreadable page yields no text rather than failing
        return blocks, text_found

    for block in text_dict.get("blocks", []):
        if block.get("type") != 0:
            continue
        block_bbox = block.get("bbox")
        if block_bbox is not None and any(
            _rect_overlap_ratio(block_bbox, region) >= _TABLE_CLAIM_OVERLAP
            for region in table_regions
        ):
            # Claimed by a table; emitting it again would duplicate the content.
            text_found = True
            continue

        line_texts: list[str] = []
        max_font_size = 0.0
        bold_chars = 0
        total_chars = 0
        is_bullet = False

        for line in block.get("lines", []):
            if _is_rotated(line):
                continue  # diagonal text is a watermark, not content
            parts: list[str] = []
            for span in line.get("spans", []):
                span_text = span.get("text", "")
                if not span_text.strip():
                    continue
                parts.append(span_text)
                size = span.get("size", 0.0)
                max_font_size = max(max_font_size, size)
                stripped = len(span_text.strip())
                total_chars += stripped
                if "bold" in (span.get("font", "") or "").lower():
                    bold_chars += stripped
            line_text = "".join(parts).strip()
            if line_text:
                line_texts.append(line_text)

        if not line_texts:
            continue

        raw_text = " ".join(line_texts)
        normalized = _normalize_chrome_text(raw_text)
        if normalized in chrome_text:
            text_found = True
            continue

        text = _demote_pua(raw_text)
        if text.lstrip().startswith(_BULLET_CHAR):
            is_bullet = True
            text = text.lstrip().lstrip(_BULLET_CHAR).strip()
        text = _join_hyphenated(text)
        if len(text) < _MIN_BLOCK_TEXT_LEN:
            continue

        text_found = True
        bbox = block_bbox or (0.0, 0.0, 0.0, 0.0)
        blocks.append(
            _ExtractedBlock(
                kind="text",
                text=text,
                font_size=max_font_size,
                is_bold=total_chars > 0 and bold_chars >= total_chars * 0.6,
                is_bullet=is_bullet,
                order=(bbox[1], bbox[0]),
            )
        )

    return blocks, text_found


def _is_rotated(line: dict[str, Any]) -> bool:
    """True when a line is not laid out horizontally."""
    direction = line.get("dir") or (1.0, 0.0)
    try:
        return abs(float(direction[1])) > 0.01
    except (TypeError, ValueError, IndexError):
        return False


def _join_hyphenated(text: str) -> str:
    """Re-join words split across a line break by a hyphen."""
    return re.sub(r"(\w)-\s+(\w)", r"\1\2", text)


def _mask_is_meaningful(doc: Any, smask_xref: int) -> bool:
    """Does this soft mask actually make anything transparent?

    A PDF writer will attach a mask that is opaque end to end - 36 of the 42 images in
    one Intel guide carry one, and not one of them hides a pixel. Attaching it anyway
    gives the image an alpha channel, which forces PNG, which is six times the bytes of
    the JPEG the page was carrying in the first place.
    """
    try:
        mask = fitz.Pixmap(doc, smask_xref)
        return min(mask.samples) < 255
    except Exception:  # policy: DEGRADED — an unreadable mask is treated as no mask
        return False


def _encode_image(
    doc: Any,
    xref: int,
    smask_xref: int,
    page_idx: int,
    warnings: list[str],
) -> tuple[bytes | None, str]:
    """Return image bytes python-docx will accept, with their extension.

    PyMuPDF hands back the stream exactly as the PDF stores it, and python-docx sniffs
    that stream for a header it knows. On a JPEG lifted out of a PDF it usually does not
    find one and raises ``UnrecognizedImageError`` - which carries no message, so the
    warning read "failed to embed image: " and forty of one document's forty-two images
    left without saying why. Rendering through a pixmap produces a header every time.

    PNG is lossless and six times larger on the screenshots these documents are made of,
    so only an image that actually carries transparency is written as PNG.
    """
    try:
        pixmap = fitz.Pixmap(doc, xref)
        if pixmap.n - pixmap.alpha >= 4:  # CMYK, separation, and other deep spaces
            pixmap = fitz.Pixmap(fitz.csRGB, pixmap)
        if smask_xref and _mask_is_meaningful(doc, smask_xref):
            # The transparency lives in a separate object; without it a logo drawn on a
            # transparent background arrives as a black rectangle.
            try:
                pixmap = fitz.Pixmap(pixmap, fitz.Pixmap(doc, smask_xref))
            except Exception:  # policy: DEGRADED — an unusable mask still leaves the image
                pass
        if pixmap.alpha:
            return pixmap.tobytes("png"), "png"
        return pixmap.tobytes("jpg", jpg_quality=_JPEG_QUALITY), "jpeg"
    except Exception as exc:  # policy: LOSSY — one unreadable image must not cost the page
        warnings.append(f"Page {page_idx + 1}: could not re-encode image: {exc!r}")
        return None, ""


def _extract_images(
    page: Any,
    page_idx: int,
    doc: Any,
    repeated_xrefs: set[int],
    warnings: list[str],
) -> list[_ExtractedBlock]:
    """Extract page images, skipping repeated logos/watermarks and tiny assets."""
    blocks: list[_ExtractedBlock] = []
    try:
        images = page.get_images(full=True)
    except Exception as exc:  # policy: LOSSY — image listing failure keeps page text
        warnings.append(f"Page {page_idx + 1}: image listing failed: {exc}")
        return blocks

    for img_info in images:
        xref = img_info[0]
        if xref in repeated_xrefs:
            continue  # logo or watermark stamped on most pages
        try:
            base_image = doc.extract_image(xref)
            if not base_image or not base_image.get("image"):
                continue
            if (
                base_image.get("width", 0) < _MIN_IMAGE_DIMENSION
                or base_image.get("height", 0) < _MIN_IMAGE_DIMENSION
            ):
                continue
            order = (0.0, 0.0)
            try:
                rects = page.get_image_rects(xref)
                if rects:
                    order = (rects[0].y0, rects[0].x0)
            except Exception:  # policy: DEGRADED — position is a nicety, not a requirement
                pass
            payload, extension = _encode_image(
                doc, xref, img_info[1] if len(img_info) > 1 else 0, page_idx, warnings
            )
            if payload is None:
                continue
            blocks.append(
                _ExtractedBlock(
                    kind="image",
                    image_bytes=payload,
                    image_ext=extension,
                    order=order,
                )
            )
        except (
            Exception
        ) as exc:  # policy: LOSSY — image extraction failure should not abort page conversion
            warnings.append(f"Page {page_idx + 1}: image extraction failed: {exc}")
    return blocks


# ---------------------------------------------------------------------------
# DOCX emission
# ---------------------------------------------------------------------------


def _is_heading_candidate(blk: _ExtractedBlock) -> bool:
    """Reject stray glyphs, bullets, and running prose as heading text.

    Font size alone is far too permissive: in a 7pt datasheet almost every
    body run clears a size-ratio threshold, which turns hundreds of ordinary
    sentences into table-of-contents entries.
    """
    if blk.is_bullet:
        return False
    stripped = blk.text.strip()
    if len(stripped) < _MIN_HEADING_CHARS or len(stripped) > _MAX_HEADING_CHARS:
        return False
    # Require at least one alphanumeric character; '®' and '•' are not headings.
    if not any(ch.isalnum() for ch in stripped):
        return False
    # Prose ends in sentence punctuation; headings almost never do.
    if stripped[-1] in ".:;,":
        return False
    # Multi-sentence text is a paragraph regardless of how large it is set.
    return stripped.count(". ") == 0


def _add_text_block(docx_doc: DocxDocument, blk: _ExtractedBlock, body_font_size: float) -> None:
    """Add a text block as a paragraph with body-relative heading detection."""
    heading_size = body_font_size * _HEADING_RATIO
    subheading_size = body_font_size * _SUBHEADING_RATIO

    if _is_heading_candidate(blk) and blk.font_size >= heading_size:
        docx_doc.add_heading(blk.text, level=1)
        return
    if _is_heading_candidate(blk) and blk.font_size >= subheading_size:
        docx_doc.add_heading(blk.text, level=2)
        return

    if blk.is_bullet:
        para = docx_doc.add_paragraph(style="List Bullet")
        run = para.add_run(blk.text)
    else:
        para = docx_doc.add_paragraph()
        run = para.add_run(blk.text)
    if blk.is_bold:
        run.bold = True


def _add_table_block(
    docx_doc: DocxDocument,
    blk: _ExtractedBlock,
    warnings: list[str],
    page_idx: int,
) -> None:
    """Emit a real DOCX table so the reader view renders rows and columns."""
    rows = blk.table_rows
    if not rows:
        return
    try:
        table = docx_doc.add_table(rows=len(rows), cols=len(rows[0]))
        try:
            table.style = "Table Grid"
        except KeyError:  # policy: DEGRADED — default style still renders a table
            pass
        for row_idx, row in enumerate(rows):
            for col_idx, value in enumerate(row):
                cell = table.cell(row_idx, col_idx)
                cell.text = value
                if row_idx == 0:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
    except Exception as exc:  # policy: LOSSY — a table that will not build must not abort the page
        warnings.append(f"Page {page_idx + 1}: failed to build table: {exc}")


def _add_image_block(
    docx_doc: DocxDocument,
    blk: _ExtractedBlock,
    warnings: list[str],
    page_idx: int,
) -> None:
    """Embed an image into the DOCX, capping width to page margins."""
    try:
        stream = io.BytesIO(blk.image_bytes)
        para = docx_doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.add_run().add_picture(stream, width=Inches(_MAX_IMG_WIDTH_IN))
    except (
        Exception
    ) as exc:  # policy: LOSSY — image embedding failure should not abort page conversion
        # `UnrecognizedImageError` stringifies to nothing, which is how this failure
        # stayed invisible; `!r` always names the class at least.
        warnings.append(f"Page {page_idx + 1}: failed to embed image: {exc!r}")


# ---------------------------------------------------------------------------
# AH-004: OCR helpers (behind FEATURE_FLAG_PDF_OCR)
# ---------------------------------------------------------------------------


def _is_ocr_enabled() -> bool:
    """Check if OCR is enabled via feature flag."""
    try:
        from app.feature_flags import BackendFeatureFlag, is_backend_feature_enabled

        return is_backend_feature_enabled(BackendFeatureFlag.PDF_OCR)
    except Exception:  # policy: DEGRADED — OCR feature-flag lookup failure disables OCR safely
        return False


def _ocr_page(page, page_idx: int, warnings: list[str]) -> str:
    """Run OCR on a rendered page image.  Returns extracted text or empty string."""
    try:
        import pytesseract
        from PIL import Image

        _OCR_DPI = 150
        _OCR_MAX_PIXELS = 25_000_000  # ~5000x5000

        pix = page.get_pixmap(dpi=_OCR_DPI)
        if pix.width * pix.height > _OCR_MAX_PIXELS:
            warnings.append(
                f"Page {page_idx + 1}: OCR skipped, page image too large ({pix.width}x{pix.height})"
            )
            return ""
        img = Image.frombytes("RGB", (pix.width, pix.height), pix.samples)
        text = pytesseract.image_to_string(img).strip()
        if text:
            logger.info("OCR extracted %d chars from page %d", len(text), page_idx + 1)
        return text
    except ImportError:
        warnings.append(f"Page {page_idx + 1}: pytesseract not installed, OCR skipped")
        return ""
    except Exception as exc:  # policy: LOSSY — OCR failure should not abort the core PDF conversion
        warnings.append(f"Page {page_idx + 1}: OCR failed: {exc}")
        return ""
