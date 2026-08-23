"""Images have to survive the PDF, and they have to arrive at a sane weight.

PyMuPDF returns an image stream exactly as the PDF stores it, and python-docx sniffs
that stream for a header it recognises. On a JPEG lifted out of a PDF it usually finds
none and raises ``UnrecognizedImageError`` - which stringifies to nothing, so the warning
read "failed to embed image: " and forty of one Intel guide's forty-two images left
without ever saying why.

Rendering through a pixmap fixes that, and then walks into the second problem: a PDF
writer routinely attaches a soft mask that hides nothing at all. Honouring it gives the
image an alpha channel, which forces PNG, which on a page of screenshots is six times
the bytes the PDF was already carrying.
"""

from __future__ import annotations

import io

import fitz
import pytest
from docx import Document
from docx.shared import Inches

from app.conversion.pdf_to_docx import _encode_image, _mask_is_meaningful


def _photo_png(width: int = 160, height: int = 120) -> bytes:
    """A pixmap that compresses the way a screenshot does, not the way a flat fill does.

    Every pixel differs from its neighbour - a gradient with noise on top. A sparser
    fixture is mostly flat, which PNG packs beautifully and JPEG handles badly, and it
    would have shown the opposite of what these documents actually do.
    """
    pixmap = fitz.Pixmap(fitz.csRGB, fitz.IRect(0, 0, width, height), False)
    seed = 12345
    for y in range(height):
        for x in range(width):
            seed = (seed * 1103515245 + 12345) & 0x7FFFFFFF
            noise = seed % 96
            pixmap.set_pixel(
                x,
                y,
                (
                    (x * 255 // width + noise) % 256,
                    (y * 255 // height + noise) % 256,
                    ((x + y) % 256 + noise) % 256,
                ),
            )
    return pixmap.tobytes("png")


def _pdf_with_image(image: bytes, mask: bytes | None = None) -> fitz.Document:
    doc = fitz.open()
    page = doc.new_page()
    page.insert_image(fitz.Rect(20, 20, 260, 180), stream=image, mask=mask)
    reopened = fitz.open("pdf", doc.tobytes())
    doc.close()
    return reopened


def _first_image(doc: fitz.Document) -> tuple[int, int]:
    """Return the (xref, smask xref) of the first image on the first page."""
    info = doc[0].get_images(full=True)[0]
    return info[0], info[1]


class TestEncodeImage:
    def test_produces_bytes_python_docx_will_accept(self):
        """The whole point: the stream that comes back can actually be embedded."""
        doc = _pdf_with_image(_photo_png())
        try:
            xref, smask = _first_image(doc)
            warnings: list[str] = []
            payload, extension = _encode_image(doc, xref, smask, 0, warnings)
        finally:
            doc.close()

        assert payload is not None
        assert warnings == []
        # python-docx is the component that used to reject these outright.
        Document().add_paragraph().add_run().add_picture(io.BytesIO(payload), width=Inches(6.0))
        assert extension in {"jpeg", "png"}

    def test_an_opaque_image_is_written_as_jpeg(self):
        """PNG is lossless and six times the size on the screenshots these documents hold."""
        doc = _pdf_with_image(_photo_png())
        try:
            xref, smask = _first_image(doc)
            payload, extension = _encode_image(doc, xref, smask, 0, [])
            png_size = len(fitz.Pixmap(doc, xref).tobytes("png"))
        finally:
            doc.close()

        assert extension == "jpeg"
        assert payload is not None
        assert len(payload) < png_size

    def test_an_unreadable_image_warns_with_something_to_read(self):
        """The failure this file exists for was invisible because the message was empty."""
        doc = _pdf_with_image(_photo_png())
        warnings: list[str] = []
        try:
            payload, _ = _encode_image(doc, 999999, 0, 4, warnings)
        finally:
            doc.close()

        assert payload is None
        assert len(warnings) == 1
        assert warnings[0].startswith("Page 5:")
        # Not "...image: " with nothing after it.
        assert warnings[0].rstrip()[-1] not in {":", " "}


class TestMaskIsMeaningful:
    def test_a_mask_that_hides_nothing_is_ignored(self):
        """36 of 42 images in one guide carried a mask; none of them hid a pixel."""
        opaque = fitz.Pixmap(fitz.csGRAY, fitz.IRect(0, 0, 240, 160), False)
        opaque.clear_with(255)
        doc = _pdf_with_image(_photo_png(), mask=opaque.tobytes("png"))
        try:
            _, smask = _first_image(doc)
            if not smask:
                pytest.skip("this PyMuPDF build stored the mask inline")
            assert _mask_is_meaningful(doc, smask) is False
        finally:
            doc.close()

    def test_a_mask_that_hides_something_is_honoured(self):
        half = fitz.Pixmap(fitz.csGRAY, fitz.IRect(0, 0, 240, 160), False)
        half.clear_with(255)
        for x in range(120):
            for y in range(160):
                half.set_pixel(x, y, (0,))
        doc = _pdf_with_image(_photo_png(), mask=half.tobytes("png"))
        try:
            _, smask = _first_image(doc)
            if not smask:
                pytest.skip("this PyMuPDF build stored the mask inline")
            assert _mask_is_meaningful(doc, smask) is True
        finally:
            doc.close()

    def test_a_mask_that_cannot_be_read_is_treated_as_no_mask(self):
        doc = _pdf_with_image(_photo_png())
        try:
            assert _mask_is_meaningful(doc, 999999) is False
        finally:
            doc.close()
