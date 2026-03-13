from __future__ import annotations

import base64
import html
import zipfile
from pathlib import Path

from docx import Document
from docx.shared import Inches
from PIL import Image, ImageDraw

SLIDE_RELATIONSHIP_TYPE = (
    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/slide"
)
IMAGE_RELATIONSHIP_TYPE = (
    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/image"
)
NOTES_RELATIONSHIP_TYPE = (
    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/notesSlide"
)
TINY_PNG_BYTES = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAEAAAABACAIAAAAlC+aJAAAAZ0lEQVR4nO3PQQ0AIBDAMMC/5+ECjiYKenVNd8f+NQJqBWgFaAVoBWgFaAVoBWgFaAVoBWgFaAVoBWgFaAVoBWgFaAVoBWgFaAVoBWgFaAVoBWgFaAVoBWgFaAVoBWgFaAW+pv0Bf6j2SFIAAAAASUVORK5CYII="
)


def _write_docx_fixture(path: Path) -> None:
    document = Document()
    document.core_properties.title = "Wave Y Rich Fixture"
    document.core_properties.author = "Codex"

    title = document.add_paragraph("Wave Y Extractor Fixture")
    title.style = "Title"

    document.add_heading("Release Overview", level=1)
    paragraph = document.add_paragraph()
    paragraph.add_run("This paragraph verifies ")
    bold_run = paragraph.add_run("bold")
    bold_run.bold = True
    paragraph.add_run(", ")
    italic_run = paragraph.add_run("italic")
    italic_run.italic = True
    paragraph.add_run(", ")
    underline_run = paragraph.add_run("underline")
    underline_run.underline = True
    paragraph.add_run(", and ")
    code_run = paragraph.add_run("build_wave_y()")
    code_run.font.name = "Consolas"
    paragraph.add_run(" extraction output.")

    document.add_paragraph("Upload DOCX through the management UI", style="List Bullet")
    document.add_paragraph("Verify semantic headings and lists", style="List Bullet 2")
    document.add_paragraph("Confirm responsive tables", style="List Bullet")
    document.add_paragraph("Check extracted images", style="List Bullet 2")
    document.add_paragraph("Open reader-ready preview", style="List Number")
    document.add_paragraph("Validate warning banner stays hidden", style="List Number 2")
    document.add_paragraph("Approve rollout", style="List Number")

    document.add_heading("Readiness Matrix", level=2)
    table = document.add_table(rows=4, cols=3)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Capability"
    table.rows[0].cells[1].text = "Owner"
    table.rows[0].cells[2].text = "Status"
    table.rows[1].cells[0].text = "DOCX extraction"
    table.rows[1].cells[1].text = "Platform"
    table.rows[1].cells[2].text = "Ready"
    table.rows[2].cells[0].text = "PPTX extraction"
    table.rows[2].cells[1].text = "Platform"
    table.rows[2].cells[2].text = "Ready"
    table.rows[3].cells[0].text = "Frontend renderer"
    table.rows[3].cells[1].text = "Web"
    table.rows[3].cells[2].text = "Ready"
    table.cell(0, 0).merge(table.cell(0, 1))
    table.cell(1, 0).merge(table.cell(2, 0))

    document.add_heading("Architecture Snapshot", level=2)
    document.add_paragraph("The embedded figure below exercises extracted image rendering.")

    image_path = path.with_name("wave_y_fixture_image.png")
    _write_fixture_png(image_path)
    try:
        document.add_picture(str(image_path), width=Inches(1.35))
    finally:
        image_path.unlink(missing_ok=True)

    document.add_paragraph("Final verification paragraph after the table and image.")
    document.save(path)


def _write_empty_docx_fixture(path: Path) -> None:
    document = Document()
    document.core_properties.title = "Wave Y Empty Fixture"
    document.core_properties.author = "Codex"
    document.save(path)


def _write_fixture_png(path: Path) -> None:
    image = Image.new("RGB", (96, 96), color=(11, 94, 144))
    draw = ImageDraw.Draw(image)
    draw.rectangle((14, 14, 82, 82), outline=(255, 255, 255), width=5)
    draw.line((14, 82, 82, 14), fill=(153, 220, 255), width=5)
    image.save(path, format="PNG")


def _write_pptx_fixture(path: Path) -> None:
    slide_entries = [
        {
            "relationship_id": "rId3",
            "target": "slides/slide1.xml",
            "xml": _make_slide(
                _make_shape(
                    10,
                    [_make_paragraph("Wave Y Launch")],
                    placeholder_type="title",
                    name="Title",
                ),
                _make_shape(
                    20,
                    [
                        _make_paragraph("Quarterly readiness review"),
                        _make_paragraph("DOCX uploads now extract cleanly", bullet=True),
                        _make_paragraph(
                            {"text": "PowerPoint decks", "bold": True},
                            " render as vertical slides",
                            bullet=True,
                        ),
                        _make_paragraph("Warnings surface only when needed", bullet=True, level=1),
                    ],
                    placeholder_type="body",
                    name="Body",
                ),
                _make_picture(
                    30,
                    "rIdImage1",
                    name="Architecture",
                    description="Architecture snapshot",
                ),
            ),
            "relationships": [
                ("rIdImage1", IMAGE_RELATIONSHIP_TYPE, "../media/wave-y-architecture.png"),
                ("rIdNotes1", NOTES_RELATIONSHIP_TYPE, "../notesSlides/notesSlide1.xml"),
            ],
        },
        {
            "relationship_id": "rId4",
            "target": "slides/slide2.xml",
            "xml": _make_slide(
                _make_shape(
                    40,
                    [_make_paragraph("Deployment Checklist")],
                    placeholder_type="title",
                    name="Title",
                ),
                _make_shape(
                    50,
                    [
                        _make_paragraph("Upload the DOCX fixture", ordered=True),
                        _make_paragraph("Verify table rendering", ordered=True),
                        _make_paragraph("Confirm image lightbox", ordered=True),
                        _make_paragraph("Inspect mobile scroll behavior", ordered=True, level=1),
                    ],
                    placeholder_type="body",
                    name="Checklist",
                ),
                _make_picture(
                    60,
                    "rIdImage2",
                    name="Checklist",
                    description="Checklist illustration",
                ),
            ),
            "relationships": [
                ("rIdImage2", IMAGE_RELATIONSHIP_TYPE, "../media/wave-y-checklist.png"),
            ],
        },
        {
            "relationship_id": "rId5",
            "target": "slides/slide3.xml",
            "xml": _make_slide(
                _make_shape(
                    70,
                    [
                        _make_paragraph("Appendix details"),
                        _make_paragraph({"text": "Fallback titles", "italic": True}, " are still usable."),
                    ],
                    name="Appendix",
                ),
            ),
        },
    ]

    slide_overrides = "".join(
        (
            '<Override PartName="/ppt/'
            f'{slide["target"]}" '
            'ContentType="application/vnd.openxmlformats-officedocument.'
            'presentationml.slide+xml"/>'
        )
        for slide in slide_entries
    )
    notes_override = (
        '<Override PartName="/ppt/notesSlides/notesSlide1.xml" '
        'ContentType="application/vnd.openxmlformats-officedocument.'
        'presentationml.notesSlide+xml"/>'
    )
    slide_id_list = "".join(
        f'<p:sldId id="{256 + index}" r:id="{slide["relationship_id"]}"/>'
        for index, slide in enumerate(slide_entries)
    )
    presentation_relationships = "".join(
        (
            f'<Relationship Id="{slide["relationship_id"]}" '
            f'Type="{SLIDE_RELATIONSHIP_TYPE}" '
            f'Target="{slide["target"]}"/>'
        )
        for slide in slide_entries
    )

    with zipfile.ZipFile(path, "w") as archive:
        archive.writestr(
            "[Content_Types].xml",
            (
                '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
                '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
                '<Default Extension="rels" '
                'ContentType="application/vnd.openxmlformats-package.relationships+xml"/>'
                '<Default Extension="xml" ContentType="application/xml"/>'
                '<Default Extension="png" ContentType="image/png"/>'
                '<Override PartName="/ppt/presentation.xml" '
                'ContentType="application/vnd.openxmlformats-officedocument.'
                'presentationml.presentation.main+xml"/>'
                '<Override PartName="/docProps/core.xml" '
                'ContentType="application/vnd.openxmlformats-package.core-properties+xml"/>'
                f"{slide_overrides}{notes_override}"
                "</Types>"
            ),
        )
        archive.writestr(
            "_rels/.rels",
            (
                '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
                '<Relationships '
                'xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
                '<Relationship Id="rId1" '
                'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/'
                'officeDocument" Target="ppt/presentation.xml"/>'
                '<Relationship Id="rId2" '
                'Type="http://schemas.openxmlformats.org/package/2006/relationships/'
                'metadata/core-properties" Target="docProps/core.xml"/>'
                "</Relationships>"
            ),
        )
        archive.writestr(
            "docProps/core.xml",
            (
                '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
                '<cp:coreProperties '
                'xmlns:cp="http://schemas.openxmlformats.org/package/2006/metadata/'
                'core-properties" '
                'xmlns:dc="http://purl.org/dc/elements/1.1/" '
                'xmlns:dcterms="http://purl.org/dc/terms/" '
                'xmlns:dcmitype="http://purl.org/dc/dcmitype/" '
                'xmlns:xsi="http://www.w3.org/2001/XMLSchema-instance">'
                "<dc:title>Wave Y Rich Deck</dc:title>"
                "<dc:creator>Codex</dc:creator>"
                '<dcterms:created xsi:type="dcterms:W3CDTF">2026-03-11T08:00:00Z'
                "</dcterms:created>"
                '<dcterms:modified xsi:type="dcterms:W3CDTF">2026-03-11T09:00:00Z'
                "</dcterms:modified>"
                "</cp:coreProperties>"
            ),
        )
        archive.writestr(
            "ppt/presentation.xml",
            (
                '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
                '<p:presentation '
                'xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" '
                'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" '
                'xmlns:p="http://schemas.openxmlformats.org/presentationml/2006/main">'
                f"<p:sldIdLst>{slide_id_list}</p:sldIdLst>"
                "</p:presentation>"
            ),
        )
        archive.writestr(
            "ppt/_rels/presentation.xml.rels",
            (
                '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
                '<Relationships '
                'xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
                f"{presentation_relationships}"
                "</Relationships>"
            ),
        )

        for slide in slide_entries:
            target = str(slide["target"])
            archive.writestr(f"ppt/{target}", str(slide["xml"]))
            archive.writestr(
                _slide_relationships_path(target),
                _relationships_xml(slide.get("relationships", [])),
            )

        archive.writestr(
            "ppt/media/wave-y-architecture.png",
            TINY_PNG_BYTES,
        )
        archive.writestr(
            "ppt/media/wave-y-checklist.png",
            TINY_PNG_BYTES,
        )
        archive.writestr(
            "ppt/notesSlides/notesSlide1.xml",
            _make_notes(
                _make_paragraph("Call out extraction confidence."),
                _make_paragraph({"text": "Mention image lightbox.", "italic": True}),
            ),
        )


def _write_empty_pptx_fixture(path: Path) -> None:
    with zipfile.ZipFile(path, "w") as archive:
        archive.writestr(
            "[Content_Types].xml",
            (
                '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
                '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
                '<Default Extension="rels" '
                'ContentType="application/vnd.openxmlformats-package.relationships+xml"/>'
                '<Default Extension="xml" ContentType="application/xml"/>'
                '<Override PartName="/ppt/presentation.xml" '
                'ContentType="application/vnd.openxmlformats-officedocument.'
                'presentationml.presentation.main+xml"/>'
                '<Override PartName="/docProps/core.xml" '
                'ContentType="application/vnd.openxmlformats-package.core-properties+xml"/>'
                "</Types>"
            ),
        )
        archive.writestr(
            "_rels/.rels",
            (
                '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
                '<Relationships '
                'xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
                '<Relationship Id="rId1" '
                'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/'
                'officeDocument" Target="ppt/presentation.xml"/>'
                '<Relationship Id="rId2" '
                'Type="http://schemas.openxmlformats.org/package/2006/relationships/'
                'metadata/core-properties" Target="docProps/core.xml"/>'
                "</Relationships>"
            ),
        )
        archive.writestr(
            "docProps/core.xml",
            (
                '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
                '<cp:coreProperties '
                'xmlns:cp="http://schemas.openxmlformats.org/package/2006/metadata/'
                'core-properties" '
                'xmlns:dc="http://purl.org/dc/elements/1.1/" '
                'xmlns:dcterms="http://purl.org/dc/terms/" '
                'xmlns:dcmitype="http://purl.org/dc/dcmitype/" '
                'xmlns:xsi="http://www.w3.org/2001/XMLSchema-instance">'
                "<dc:title>Wave Y Empty Deck</dc:title>"
                "<dc:creator>Codex</dc:creator>"
                "</cp:coreProperties>"
            ),
        )
        archive.writestr(
            "ppt/presentation.xml",
            (
                '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
                '<p:presentation '
                'xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" '
                'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" '
                'xmlns:p="http://schemas.openxmlformats.org/presentationml/2006/main">'
                "<p:sldIdLst></p:sldIdLst>"
                "</p:presentation>"
            ),
        )
        archive.writestr(
            "ppt/_rels/presentation.xml.rels",
            (
                '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
                '<Relationships '
                'xmlns="http://schemas.openxmlformats.org/package/2006/relationships"></Relationships>'
            ),
        )


def _slide_relationships_path(target: str) -> str:
    slide_name = Path(target).name
    return f"ppt/slides/_rels/{slide_name}.rels"


def _relationships_xml(relationships: list[tuple[str, str, str]]) -> str:
    items = "".join(
        (
            f'<Relationship Id="{relationship_id}" '
            f'Type="{relationship_type}" '
            f'Target="{target}"/>'
        )
        for relationship_id, relationship_type, target in relationships
    )
    return (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<Relationships '
        'xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
        f"{items}"
        "</Relationships>"
    )


def _make_slide(*nodes: str) -> str:
    return (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<p:sld '
        'xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" '
        'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" '
        'xmlns:p="http://schemas.openxmlformats.org/presentationml/2006/main" '
        'xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture">'
        '<p:cSld><p:spTree>'
        f'{"".join(nodes)}'
        "</p:spTree></p:cSld>"
        "</p:sld>"
    )


def _make_notes(*paragraphs: str) -> str:
    return (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<p:notes '
        'xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" '
        'xmlns:p="http://schemas.openxmlformats.org/presentationml/2006/main">'
        '<p:cSld><p:spTree>'
        f'{_make_shape(901, list(paragraphs), placeholder_type="body", name="Notes")}'
        "</p:spTree></p:cSld>"
        "</p:notes>"
    )


def _make_shape(
    shape_id: int,
    paragraphs: list[str],
    *,
    placeholder_type: str | None = None,
    name: str | None = None,
) -> str:
    placeholder_xml = f'<p:ph type="{placeholder_type}"/>' if placeholder_type else ""
    return (
        "<p:sp>"
        "<p:nvSpPr>"
        f'<p:cNvPr id="{shape_id}" name="{html.escape(name or f"Shape {shape_id}")}"/>'
        "<p:cNvSpPr/>"
        f"<p:nvPr>{placeholder_xml}</p:nvPr>"
        "</p:nvSpPr>"
        "<p:spPr/>"
        "<p:txBody><a:bodyPr/><a:lstStyle/>"
        f'{"".join(paragraphs)}'
        "</p:txBody>"
        "</p:sp>"
    )


def _make_picture(
    picture_id: int,
    relationship_id: str,
    *,
    name: str,
    description: str | None = None,
) -> str:
    description_attr = f' descr="{html.escape(description)}"' if description else ""
    return (
        "<p:pic>"
        "<p:nvPicPr>"
        f'<p:cNvPr id="{picture_id}" name="{html.escape(name)}"{description_attr}/>'
        "<p:cNvPicPr/>"
        "<p:nvPr/>"
        "</p:nvPicPr>"
        "<p:blipFill>"
        f'<a:blip r:embed="{relationship_id}"/>'
        "<a:stretch><a:fillRect/></a:stretch>"
        "</p:blipFill>"
        "<p:spPr/>"
        "</p:pic>"
    )


def _make_paragraph(
    *runs: str | dict[str, object],
    bullet: bool = False,
    ordered: bool = False,
    level: int = 0,
) -> str:
    p_pr = ""
    if bullet or ordered or level:
        bullet_xml = ""
        if ordered:
            bullet_xml = '<a:buAutoNum type="arabicPeriod" startAt="1"/>'
        elif bullet:
            bullet_xml = '<a:buChar char="•"/>'
        level_attr = f' lvl="{level}"' if level else ""
        p_pr = f"<a:pPr{level_attr}>{bullet_xml}</a:pPr>"
    return f"<a:p>{p_pr}{''.join(_make_run(run) for run in runs)}</a:p>"


def _make_run(run: str | dict[str, object]) -> str:
    if isinstance(run, str):
        return f"<a:r><a:t>{html.escape(run)}</a:t></a:r>"

    attributes: list[str] = []
    if run.get("bold"):
        attributes.append(' b="1"')
    if run.get("italic"):
        attributes.append(' i="1"')

    return (
        f'<a:r><a:rPr{"".join(attributes)}/>'
        f'<a:t>{html.escape(str(run["text"]))}</a:t></a:r>'
    )


def main() -> None:
    fixture_dir = Path(__file__).resolve().parent
    fixture_dir.mkdir(parents=True, exist_ok=True)

    _write_docx_fixture(fixture_dir / "wave_y_rich.docx")
    _write_empty_docx_fixture(fixture_dir / "wave_y_empty.docx")
    _write_pptx_fixture(fixture_dir / "wave_y_rich.pptx")
    _write_empty_pptx_fixture(fixture_dir / "wave_y_empty.pptx")


if __name__ == "__main__":
    main()
