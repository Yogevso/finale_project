from __future__ import annotations

import zipfile
from collections.abc import Callable
from pathlib import Path

from docx import Document
from PIL import Image

from app.conversion.docx_extractor import extract_docx
from app.conversion.html_generator import ir_to_html

FIXTURE_DIR = Path(__file__).resolve().parent / "fixtures" / "documents"
RICH_DOCX_FIXTURE = FIXTURE_DIR / "wave_y_rich.docx"
EMPTY_DOCX_FIXTURE = FIXTURE_DIR / "wave_y_empty.docx"


def _write_docx(path: Path, *, title: str | None = None, author: str | None = None) -> Document:
    document = Document()
    if title is not None:
        document.core_properties.title = title
    if author is not None:
        document.core_properties.author = author
    document.save(path)
    return document


def _write_png(path: Path, *, color: tuple[int, int, int] = (24, 86, 120)) -> None:
    image = Image.new("RGB", (24, 24), color=color)
    image.save(path, format="PNG")


def _rewrite_docx_without_media(source_path: Path, target_path: Path) -> None:
    with zipfile.ZipFile(source_path) as source_archive:
        media_member = next(
            name for name in source_archive.namelist() if name.startswith("word/media/")
        )
        with zipfile.ZipFile(target_path, "w") as target_archive:
            for member in source_archive.infolist():
                if member.filename == media_member:
                    continue
                target_archive.writestr(member, source_archive.read(member.filename))


def _rewrite_docx_document_xml(
    source_path: Path,
    target_path: Path,
    *,
    transform: Callable[[str], str],
) -> None:
    with zipfile.ZipFile(source_path) as source_archive:
        document_xml = source_archive.read("word/document.xml").decode("utf-8")
        updated_xml = transform(document_xml)
        with zipfile.ZipFile(target_path, "w") as target_archive:
            for member in source_archive.infolist():
                if member.filename == "word/document.xml":
                    target_archive.writestr(member, updated_xml.encode("utf-8"))
                    continue
                target_archive.writestr(member, source_archive.read(member.filename))


def test_extract_docx_returns_ready_result_for_simple_document(tmp_path):
    docx_path = tmp_path / "simple.docx"
    document = _write_docx(docx_path, title="Wave Y Spec", author="Codex")
    document.add_paragraph("Paragraph extraction works.")
    document.save(docx_path)

    result = extract_docx(docx_path)

    assert result.status == "ready"
    assert result.title == "Wave Y Spec"
    assert result.metadata["title"] == "Wave Y Spec"
    assert result.metadata["author"] == "Codex"
    assert "<article class=\"docx-document\"" in result.html
    assert '<p class="extracted-paragraph">Paragraph extraction works.</p>' in result.html
    assert result.ir is not None
    assert ir_to_html(result.ir) == result.html
    assert result.warnings == []
    assert result.confidence == 1.0
    assert result.extraction_error is None


def test_extract_docx_renders_headings_from_word_styles(tmp_path):
    docx_path = tmp_path / "headings.docx"
    document = Document()

    title_paragraph = document.add_paragraph("Project Atlas")
    title_paragraph.style = "Title"
    document.add_heading("Introduction", level=1)
    document.add_heading("Implementation Details", level=2)
    document.add_heading("Implementation Details", level=2)
    document.save(docx_path)

    result = extract_docx(docx_path)

    assert result.status == "ready"
    assert result.title == "Project Atlas"
    assert (
        '<h1 class="extracted-heading extracted-heading-level-1" '
        'id="heading-project-atlas">Project Atlas</h1>'
    ) in result.html
    assert (
        '<h1 class="extracted-heading extracted-heading-level-1" '
        'id="heading-introduction">Introduction</h1>'
    ) in result.html
    assert (
        '<h2 class="extracted-heading extracted-heading-level-2" '
        'id="heading-implementation-details">Implementation Details</h2>'
        in result.html
    )
    assert (
        '<h2 class="extracted-heading extracted-heading-level-2" '
        'id="heading-implementation-details-2">Implementation Details</h2>'
        in result.html
    )
    assert [(heading.id, heading.level, heading.text) for heading in result.headings] == [
        ("heading-project-atlas", 1, "Project Atlas"),
        ("heading-introduction", 1, "Introduction"),
        ("heading-implementation-details", 2, "Implementation Details"),
        ("heading-implementation-details-2", 2, "Implementation Details"),
    ]


def test_extract_docx_renders_inline_run_formatting(tmp_path):
    docx_path = tmp_path / "inline-formatting.docx"
    document = Document()

    heading = document.add_paragraph(style="Heading 1")
    heading.add_run("Inline ")
    heading_code = heading.add_run("Code")
    heading_code.font.name = "Consolas"

    paragraph = document.add_paragraph()
    paragraph.add_run("Use ")
    bold_run = paragraph.add_run("bold")
    bold_run.bold = True
    paragraph.add_run(", ")
    italic_run = paragraph.add_run("italic")
    italic_run.italic = True
    paragraph.add_run(", ")
    underline_run = paragraph.add_run("underline")
    underline_run.underline = True
    paragraph.add_run(", ")
    combo_run = paragraph.add_run("combo")
    combo_run.bold = True
    combo_run.italic = True
    paragraph.add_run(", and ")
    code_run = paragraph.add_run("print('hi')")
    code_run.font.name = "Consolas"
    paragraph.add_run(".")
    document.save(docx_path)

    result = extract_docx(docx_path)

    assert result.status == "ready"
    assert (
        '<h1 class="extracted-heading extracted-heading-level-1" '
        'id="heading-inline-code">Inline <code class="extracted-code">Code</code></h1>'
    ) in result.html
    assert (
        '<p class="extracted-paragraph">Use <strong>bold</strong>, <em>italic</em>, '
        '<u>underline</u>, <strong><em>combo</em></strong>, and '
        '<code class="extracted-code">print(&#x27;hi&#x27;)</code>.</p>'
    ) in result.html


def test_extract_docx_groups_bulleted_and_numbered_lists_with_nesting(tmp_path):
    docx_path = tmp_path / "lists.docx"
    document = Document()

    document.add_paragraph("First bullet", style="List Bullet")
    document.add_paragraph("Nested bullet", style="List Bullet 2")
    document.add_paragraph("Second bullet", style="List Bullet")
    document.add_paragraph("First number", style="List Number")
    document.add_paragraph("Nested number", style="List Number 2")
    document.add_paragraph("Second number", style="List Number")
    document.add_paragraph("Gap paragraph.")
    document.add_paragraph("Restarted bullet", style="List Bullet")
    document.save(docx_path)

    result = extract_docx(docx_path)

    assert result.status == "ready"
    assert (
        '<ul class="extracted-list"><li>First bullet<ul><li>Nested bullet</li></ul></li>'
        "<li>Second bullet</li></ul>"
    ) in result.html
    assert (
        '<ol class="extracted-list"><li>First number<ol><li>Nested number</li></ol></li>'
        "<li>Second number</li></ol>"
    ) in result.html
    assert '<p class="extracted-paragraph">Gap paragraph.</p>' in result.html
    assert (
        '<p class="extracted-paragraph">Gap paragraph.</p>'
        '<ul class="extracted-list"><li>Restarted bullet</li></ul>'
    ) in result.html


def test_extract_docx_renders_tables_with_header_body_and_merged_cells(tmp_path):
    docx_path = tmp_path / "tables.docx"
    document = Document()
    document.add_paragraph("Before table.")

    table = document.add_table(rows=3, cols=3)
    table.rows[0].cells[0].text = "Header 1"
    table.rows[0].cells[1].text = "Header 2"
    table.rows[0].cells[2].text = "Header 3"
    table.rows[1].cells[0].text = "Cell 1"
    table.rows[1].cells[1].text = "Cell 2"
    table.rows[1].cells[2].text = "Cell 3"
    table.rows[2].cells[0].text = "Cell 4"
    table.rows[2].cells[1].text = "Cell 5"
    table.rows[2].cells[2].text = "Cell 6"
    table.cell(0, 0).merge(table.cell(0, 1))
    table.cell(1, 0).merge(table.cell(2, 0))

    document.add_paragraph("After table.")
    document.save(docx_path)

    result = extract_docx(docx_path)

    assert result.status == "ready"
    assert (
        '<p class="extracted-paragraph">Before table.</p>'
        '<div class="table-wrapper"><table class="extracted-table">'
        in result.html
    )
    assert (
        '<thead><tr><th colspan="2"><p class="extracted-paragraph">Header 1</p>'
        '<p class="extracted-paragraph">Header 2</p></th>'
    ) in result.html
    assert '<th><p class="extracted-paragraph">Header 3</p></th></tr></thead>' in result.html
    assert (
        '<tbody><tr><td rowspan="2"><p class="extracted-paragraph">Cell 1</p>'
        '<p class="extracted-paragraph">Cell 4</p></td>'
        '<td><p class="extracted-paragraph">Cell 2</p></td>'
        '<td><p class="extracted-paragraph">Cell 3</p></td></tr>'
        '<tr><td><p class="extracted-paragraph">Cell 5</p></td>'
        '<td><p class="extracted-paragraph">Cell 6</p></td></tr></tbody>'
    ) in result.html
    assert '</table></div><p class="extracted-paragraph">After table.</p>' in result.html
    assert [warning.code for warning in result.warnings] == []


def test_extract_docx_renders_embedded_images_as_base64_figures_in_document_order(tmp_path):
    image_path = tmp_path / "diagram.png"
    _write_png(image_path)

    docx_path = tmp_path / "with-image.docx"
    document = Document()
    document.add_paragraph("Before image.")
    document.add_picture(str(image_path))
    document.add_paragraph("After image.")
    document.save(docx_path)

    result = extract_docx(docx_path)

    assert result.status == "ready"
    assert (
        '<p class="extracted-paragraph">Before image.</p>'
        '<figure class="extracted-image"><img '
        "src=\"data:image/png;base64,"
    ) in result.html
    assert (
        'alt="Figure 1" loading="lazy" /><figcaption class="extracted-image-caption">'
        'Figure 1</figcaption></figure><p class="extracted-paragraph">After image.</p>'
    ) in result.html
    assert [warning.code for warning in result.warnings] == []


def test_extract_docx_inserts_placeholder_and_warning_for_missing_images(tmp_path):
    image_path = tmp_path / "missing.png"
    _write_png(image_path, color=(180, 64, 32))

    valid_path = tmp_path / "valid-image.docx"
    document = Document()
    document.add_picture(str(image_path))
    document.save(valid_path)

    broken_path = tmp_path / "broken-image.docx"
    _rewrite_docx_without_media(valid_path, broken_path)

    result = extract_docx(broken_path)

    assert result.status == "ready"
    assert (
        '<div class="extracted-image-placeholder">[Image unavailable: Figure 1]</div>'
        '<figcaption class="extracted-image-caption">Figure 1</figcaption>'
    ) in result.html
    assert [(warning.code, warning.count) for warning in result.warnings] == [
        ("MISSING_IMAGES", 1),
    ]
    assert result.confidence == 0.9


def test_extract_docx_ignores_watermark_textboxes_and_vml_textpath_shapes(tmp_path):
    base_path = tmp_path / "watermark-base.docx"
    document = Document()
    document.add_paragraph("Actual body content.")
    document.save(base_path)

    watermark_path = tmp_path / "watermark-shapes.docx"
    watermark_markup = """
<w:p>
  <w:r>
    <w:drawing>
      <wp:anchor distT="0" distB="0" distL="0" distR="0">
        <wp:docPr id="1" name="Textbox 1"/>
      </wp:anchor>
    </w:drawing>
  </w:r>
</w:p>
<w:p>
  <w:r>
    <w:pict>
      <v:shape style="rotation:315" type="#_x0000_t136">
        <v:textpath string="CONFIDENTIAL WATERMARK"/>
      </v:shape>
    </w:pict>
  </w:r>
</w:p>
"""
    _rewrite_docx_document_xml(
        base_path,
        watermark_path,
        transform=lambda xml: xml.replace("<w:sectPr", f"{watermark_markup}<w:sectPr", 1),
    )

    result = extract_docx(watermark_path)

    assert result.status == "ready"
    assert '<p class="extracted-paragraph">Actual body content.</p>' in result.html
    assert "[Image unavailable:" not in result.html
    assert "<figure class=\"extracted-image\">" not in result.html
    assert [warning.code for warning in result.warnings] == []
    assert result.confidence == 1.0


def test_extract_docx_marks_empty_documents_with_no_content_warning(tmp_path):
    docx_path = tmp_path / "empty.docx"
    _write_docx(docx_path)

    result = extract_docx(docx_path)

    assert result.status == "ready"
    assert result.html == (
        '<article class="docx-document" role="article" aria-label="Uploaded document"></article>'
    )
    assert [warning.code for warning in result.warnings] == ["NO_CONTENT"]
    assert result.confidence == 0.5


def test_extract_docx_fails_for_invalid_archives(tmp_path):
    invalid_path = tmp_path / "broken.docx"
    invalid_path.write_text("not a docx", encoding="utf-8")

    result = extract_docx(invalid_path)

    assert result.status == "failed"
    assert result.html == ""
    assert result.confidence == 0.0
    assert result.extraction_error is not None
    assert "Invalid DOCX archive" in result.extraction_error
    assert [warning.code for warning in result.warnings] == ["PARSE_FAILED"]


def test_extract_docx_renders_wave_y_rich_fixture_file():
    result = extract_docx(RICH_DOCX_FIXTURE)

    assert result.status == "ready"
    assert result.metadata["title"] == "Wave Y Rich Fixture"
    assert result.metadata["author"] == "Codex"
    assert result.ir is not None
    assert ir_to_html(result.ir) == result.html
    assert [heading.text for heading in result.headings] == [
        "Wave Y Extractor Fixture",
        "Release Overview",
        "Readiness Matrix",
        "Architecture Snapshot",
    ]
    assert (
        '<article class="docx-document" role="article" aria-label="Uploaded document">'
        in result.html
    )
    assert (
        '<h1 class="extracted-heading extracted-heading-level-1" '
        'id="heading-wave-y-extractor-fixture">Wave Y Extractor Fixture</h1>'
    ) in result.html
    assert "<strong>bold</strong>" in result.html
    assert "<em>italic</em>" in result.html
    assert "<u>underline</u>" in result.html
    assert '<code class="extracted-code">build_wave_y()</code>' in result.html
    assert (
        '<ul class="extracted-list"><li>Upload DOCX through the management UI<ul>'
        "<li>Verify semantic headings and lists</li></ul></li>"
    ) in result.html
    assert (
        '<ol class="extracted-list"><li>Open reader-ready preview<ol>'
        "<li>Validate warning banner stays hidden</li></ol></li>"
        "<li>Approve rollout</li></ol>"
    ) in result.html
    assert '<div class="table-wrapper"><table class="extracted-table">' in result.html
    assert '<th colspan="2"><p class="extracted-paragraph">Capability</p>' in result.html
    assert '<td rowspan="2"><p class="extracted-paragraph">DOCX extraction</p>' in result.html
    assert '<figure class="extracted-image"><img src="data:image/png;base64,' in result.html
    assert "Figure 1" in result.html
    assert result.warnings == []
    assert result.confidence == 1.0


def test_extract_docx_marks_wave_y_empty_fixture_with_no_content_warning():
    result = extract_docx(EMPTY_DOCX_FIXTURE)

    assert result.status == "ready"
    assert result.title == "Wave Y Empty Fixture"
    assert result.metadata["title"] == "Wave Y Empty Fixture"
    assert result.html == (
        '<article class="docx-document" role="article" aria-label="Uploaded document"></article>'
    )
    assert [warning.code for warning in result.warnings] == ["NO_CONTENT"]
    assert result.confidence == 0.5
