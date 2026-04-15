"""Conversion benchmark scenarios for the production perf gate."""

from __future__ import annotations

import json
from io import BytesIO
from time import perf_counter

import fitz
import pytest
from docx import Document as DocxDocument

from app.conversion import get_document_conversion_pipeline
from app.services.pdf_export_service import render_html_to_pdf

BENCHMARK_ITERATIONS = 5
DOCX_MIME_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def _percentiles(latencies_ms: list[float]) -> tuple[float, float]:
    ordered = sorted(latencies_ms)
    if not ordered:
        return (0.0, 0.0)

    midpoint = len(ordered) // 2
    if len(ordered) % 2 == 0:
        p50 = (ordered[midpoint - 1] + ordered[midpoint]) / 2
    else:
        p50 = ordered[midpoint]
    p95_index = max(0, min(len(ordered) - 1, int(len(ordered) * 0.95) - 1))
    return p50, ordered[p95_index]


def _measure_latency_ms(action, *, iterations: int) -> list[float]:
    samples: list[float] = []
    for _ in range(iterations):
        started_at = perf_counter()
        action()
        elapsed_ms = (perf_counter() - started_at) * 1000.0
        samples.append(elapsed_ms)
    return samples


def _build_docx_bytes() -> bytes:
    document = DocxDocument()
    document.add_heading("Performance Benchmark Document", level=1)
    document.add_paragraph(
        "This DOCX fixture exercises the structured reader-artifact conversion path."
    )
    document.add_paragraph("It is intentionally small so the benchmark is repeatable inside CI.")
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _build_pdf_bytes() -> bytes:
    document = fitz.open()
    page = document.new_page()
    page.insert_text((72, 72), "Performance Benchmark PDF", fontsize=16)
    page.insert_text(
        (72, 110), "This PDF exercises the PDF export and conversion path.", fontsize=12
    )
    pdf_bytes = document.write()
    document.close()
    return pdf_bytes


@pytest.mark.slow
@pytest.mark.integration
def test_conversion_path_benchmarks(record_property):
    """Measure representative reader-artifact and PDF-export paths."""
    pipeline = get_document_conversion_pipeline()
    docx_bytes = _build_docx_bytes()
    pdf_bytes = _build_pdf_bytes()
    html = (
        "<h1>Production Perf Gate</h1>"
        "<p>This HTML payload exercises the PDF export renderer.</p>"
    )

    docx_reader_samples = _measure_latency_ms(
        lambda: _assert_reader_artifact_ready(
            pipeline.convert_document_to_reader_artifact(
                docx_bytes,
                DOCX_MIME_TYPE,
                filename="perf-gate.docx",
            )
        ),
        iterations=BENCHMARK_ITERATIONS,
    )
    pdf_export_samples = _measure_latency_ms(
        lambda: _assert_pdf_bytes(render_html_to_pdf(html, title="Perf Gate")),
        iterations=BENCHMARK_ITERATIONS,
    )
    pdf_reader_samples = _measure_latency_ms(
        lambda: _assert_reader_artifact_ready(
            pipeline.convert_document_to_reader_artifact(
                pdf_bytes,
                "application/pdf",
                filename="perf-gate.pdf",
            )
        ),
        iterations=BENCHMARK_ITERATIONS,
    )

    docx_reader_p50, docx_reader_p95 = _percentiles(docx_reader_samples)
    pdf_export_p50, pdf_export_p95 = _percentiles(pdf_export_samples)
    pdf_reader_p50, pdf_reader_p95 = _percentiles(pdf_reader_samples)

    metrics = {
        "docx_reader_artifact": {"p50_ms": docx_reader_p50, "p95_ms": docx_reader_p95},
        "pdf_export": {"p50_ms": pdf_export_p50, "p95_ms": pdf_export_p95},
        "pdf_to_reader_artifact": {"p50_ms": pdf_reader_p50, "p95_ms": pdf_reader_p95},
    }
    record_property("conversion_benchmark_metrics_json", json.dumps(metrics, sort_keys=True))

    for metric in metrics.values():
        assert metric["p50_ms"] >= 0.0
        assert metric["p95_ms"] >= metric["p50_ms"]


def _assert_reader_artifact_ready(artifact: dict[str, object] | None) -> None:
    assert artifact is not None
    assert artifact["status"] == "ready"
    assert isinstance(artifact.get("html_content"), str)
    assert artifact["html_content"]


def _assert_pdf_bytes(pdf_bytes: bytes) -> None:
    assert pdf_bytes.startswith(b"%PDF")
    assert len(pdf_bytes) > 100
